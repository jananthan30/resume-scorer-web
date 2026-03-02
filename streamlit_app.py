"""
Resume Scorer — Consumer Web App (Streamlit)

Polished single-page app for scoring resumes against job descriptions.
Features visual score breakdowns (Plotly gauges), auth dialogs, usage dashboard,
and Stripe billing integration.

Run locally:
    streamlit run cloud/streamlit_app.py

Deploy to Streamlit Cloud:
    1. Push to repo (jananthan30/resume-scorer-web)
    2. Connect to Streamlit Cloud (share.streamlit.io)
    3. Set secrets: SCORER_API_URL, STRIPE_PUBLISHABLE_KEY
"""

import base64
import json
import os
import threading
import time
import uuid

import plotly.graph_objects as go
import requests
import streamlit as st

# ─── Configuration ───────────────────────────────────────────────────────────
API_URL = os.getenv("SCORER_API_URL", "https://resume-scorer.fly.dev")

# ─── Page config (must be first Streamlit call) ─────────────────────────────
st.set_page_config(
    page_title="Resume Scorer",
    page_icon="📄",
    layout="wide",
    initial_sidebar_state="collapsed",
)

# ─── Session state defaults ─────────────────────────────────────────────────
_defaults = {
    "token": None,
    "user": None,
    "session_id": str(uuid.uuid4()),  # Anonymous fingerprint per session
    "page": "home",
    "score_result": None,
    "scores_used": 0,
    # Cross-page handoff from Discover → Score/Rewrite
    "prefill_resume": "",
    "prefill_jd": "",
    "prefill_job_url": "",
    "prefill_job_title": "",
    "prefill_company": "",
    "discover_results": None,
    "stored_resume": "",  # Saved resume text persists across pages
    "resume_on_file": False,  # True when resume is persisted in cloud DB
}
for key, val in _defaults.items():
    if key not in st.session_state:
        st.session_state[key] = val

def _decode_jwt_payload(token: str) -> dict:
    """Decode JWT claims without signature check (server validates on every API call)."""
    try:
        payload_b64 = token.split(".")[1]
        payload_b64 += "=" * (4 - len(payload_b64) % 4)
        return json.loads(base64.b64decode(payload_b64).decode("utf-8"))
    except Exception:
        return {}


# ─── Restore session from URL param (survives page refresh) ──────────────────
if st.session_state.token is None:
    _url_token = st.query_params.get("_t", "")
    if _url_token:
        _claims = _decode_jwt_payload(_url_token)
        if _claims and _claims.get("exp", 0) > time.time():
            st.session_state.token = _url_token
            st.session_state.user = {
                "email": _claims.get("email", ""),
                "tier": _claims.get("tier", "free"),
                "user_id": _claims.get("sub", ""),
            }


# ─── Custom CSS ──────────────────────────────────────────────────────────────
st.markdown("""
<style>
    /* Dark theme */
    .stApp { background-color: #0f172a; }
    [data-testid="stHeader"] { background-color: #0f172a; }

    /* Hide Streamlit branding and header */
    #MainMenu { visibility: hidden; }
    footer { visibility: hidden; }
    [data-testid="stToolbar"] { display: none; }
    [data-testid="stHeader"] { display: none; }

    /* Card styles */
    .card {
        background: #1e293b;
        border-radius: 12px;
        padding: 24px;
        border: 1px solid #334155;
        margin-bottom: 16px;
    }
    .card-accent {
        background: #1e293b;
        border-radius: 12px;
        padding: 24px;
        border: 1px solid #6366f1;
        margin-bottom: 16px;
    }

    /* Hero stat cards */
    .stat-card {
        background: #1e293b;
        border-radius: 12px;
        padding: 20px 16px;
        text-align: center;
        border: 1px solid #334155;
    }
    .stat-num { font-size: 32px; font-weight: 700; color: #818cf8; }
    .stat-label { font-size: 12px; color: #94a3b8; margin-top: 4px; text-transform: uppercase; letter-spacing: 1px; }

    /* Keyword chips */
    .chip-matched {
        display: inline-block;
        background: #166534;
        color: #bbf7d0;
        border-radius: 20px;
        padding: 4px 12px;
        margin: 3px;
        font-size: 13px;
        font-weight: 500;
    }
    .chip-missing {
        display: inline-block;
        background: #991b1b;
        color: #fecaca;
        border-radius: 20px;
        padding: 4px 12px;
        margin: 3px;
        font-size: 13px;
        font-weight: 500;
    }

    /* Recommendation badge */
    .badge {
        display: inline-block;
        border-radius: 8px;
        padding: 6px 16px;
        font-weight: 700;
        font-size: 14px;
        text-transform: uppercase;
        letter-spacing: 0.5px;
    }
    .badge-green { background: #166534; color: #bbf7d0; }
    .badge-yellow { background: #854d0e; color: #fef08a; }
    .badge-red { background: #991b1b; color: #fecaca; }

    /* Nav pills */
    .nav-container {
        display: flex;
        gap: 8px;
        margin-bottom: 24px;
        padding: 8px 0;
        border-bottom: 1px solid #334155;
    }
    .nav-pill {
        background: transparent;
        color: #94a3b8;
        border: 1px solid #334155;
        border-radius: 8px;
        padding: 8px 20px;
        cursor: pointer;
        font-size: 14px;
        text-decoration: none;
        transition: all 0.2s;
    }
    .nav-pill:hover { background: #1e293b; color: #e2e8f0; }
    .nav-pill-active {
        background: #6366f1 !important;
        color: white !important;
        border-color: #6366f1 !important;
    }

    /* Section headers */
    .section-title {
        color: #e2e8f0;
        font-size: 28px;
        font-weight: 700;
        margin-bottom: 8px;
    }
    .section-subtitle {
        color: #94a3b8;
        font-size: 16px;
        margin-bottom: 24px;
    }

    /* Pricing */
    .price-tag {
        font-size: 40px;
        font-weight: 800;
        color: #818cf8;
    }
    .price-period { font-size: 16px; color: #94a3b8; font-weight: 400; }

    /* Progress bar override */
    .stProgress > div > div > div { background-color: #6366f1; }

    /* Remove extra padding on wide layout */
    .block-container { padding-top: 1rem; }
</style>
""", unsafe_allow_html=True)


# ─── API helpers ─────────────────────────────────────────────────────────────

def api(method: str, endpoint: str, json_data: dict = None, token: str = None) -> dict:
    """Call the scorer API with session fingerprinting."""
    url = f"{API_URL.rstrip('/')}{endpoint}"
    headers = {
        "Content-Type": "application/json",
        "X-Client-Fingerprint": st.session_state.session_id,
    }
    if token:
        headers["Authorization"] = f"Bearer {token}"
    # Long-running endpoints need generous timeouts
    if endpoint == "/jobs/fetch-jd":
        timeout = 90
    elif endpoint in ("/rewrite", "/cover-letter", "/jobs/discover"):
        timeout = 150
    elif method == "GET":
        timeout = 30
    else:
        timeout = 60
    try:
        if method == "GET":
            r = requests.get(url, headers=headers, timeout=timeout)
        elif method == "DELETE":
            r = requests.delete(url, headers=headers, timeout=timeout)
        elif method == "PUT":
            r = requests.put(url, json=json_data or {}, headers=headers, timeout=timeout)
        else:
            r = requests.post(url, json=json_data or {}, headers=headers, timeout=timeout)
        try:
            data = r.json()
        except ValueError:
            # Empty body or HTML error page (proxy timeout, 502, etc.)
            data = {"detail": r.text[:300].strip() if r.text.strip() else f"Empty response (HTTP {r.status_code})"}
        return {"status": r.status_code, "data": data}
    except requests.RequestException as e:
        return {"status": 0, "data": {"detail": str(e)}}


def is_authenticated() -> bool:
    return st.session_state.token is not None and st.session_state.user is not None


# ─── Cloud resume storage helpers ────────────────────────────────────────────

def _fetch_saved_resume(token: str) -> str:
    """Fetch the user's saved resume from cloud DB. Returns text or ''."""
    result = api("GET", "/resume", token=token)
    if result["status"] == 200:
        return result["data"].get("resume_text", "")
    return ""


def _upload_resume_to_cloud(token: str, text: str, filename: str = "resume.txt") -> bool:
    """Push resume text to cloud DB silently. Returns True on success."""
    result = api("POST", "/resume/upload", {"resume_text": text, "filename": filename}, token=token)
    return result["status"] == 200


def _delete_resume_from_cloud(token: str) -> bool:
    """Delete the saved resume from cloud DB. Returns True on success."""
    result = api("DELETE", "/resume", token=token)
    return result["status"] == 200


def api_stream(endpoint: str, json_data: dict, token: str, timeout: int = 180):
    """
    POST to an SSE endpoint and yield parsed event dicts as they arrive.
    Handles non-streaming fallback (plain JSON response) transparently.
    Yields dicts; callers check event["stage"] for "done" / "error" / progress.
    """
    url = f"{API_URL.rstrip('/')}{endpoint}"
    headers = {
        "Content-Type": "application/json",
        "Accept": "text/event-stream",
        "X-Client-Fingerprint": st.session_state.session_id,
    }
    if token:
        headers["Authorization"] = f"Bearer {token}"
    try:
        with requests.post(url, json=json_data, headers=headers, stream=True, timeout=timeout) as r:
            if r.status_code != 200:
                try:
                    detail = r.json().get("detail", f"HTTP {r.status_code}")
                except Exception:
                    detail = r.text[:300] or f"HTTP {r.status_code}"
                yield {"stage": "error", "detail": detail, "status": r.status_code}
                return
            for raw_line in r.iter_lines():
                if raw_line and raw_line.startswith(b"data: "):
                    try:
                        yield json.loads(raw_line[6:])
                    except json.JSONDecodeError:
                        pass
    except Exception as exc:
        yield {"stage": "error", "detail": str(exc)}


def _fetch_full_jd(job_url: str, job_title: str, token: str) -> tuple:
    """Scrape full JD from listing URL. Returns (jd_text, error_message)."""
    result = api(
        "POST",
        "/jobs/fetch-jd",
        {"url": job_url, "job_title": job_title, "use_ai": True},
        token=token,
    )
    if result["status"] == 200:
        return result["data"].get("jd_text", ""), ""
    return "", result["data"].get("detail", "Could not fetch the job description.")


def _jd_is_truncated(text: str) -> bool:
    """Return True when the description looks like an API snippet rather than a full JD."""
    t = text.strip()
    # Explicit truncation markers used by Adzuna / Remotive APIs
    if t.endswith("…") or t.endswith("...") or t.endswith("…"):
        return True
    # Short text is almost always a snippet
    if len(t) < 1200:
        return True
    return False


def _apply_jd_prefill(jd_key: str, prefill_jd: str, job_url: str, job_title: str):
    """
    Called once per navigation from a job card.
    If the description looks truncated, scrapes the listing URL for the full text.
    Falls back to the API snippet silently on failure.
    Returns (final_jd_text, warning_message).
    """
    if _jd_is_truncated(prefill_jd) and job_url and is_authenticated():
        with st.spinner("Fetching full job description from listing…"):
            full_jd, err = _fetch_full_jd(job_url, job_title, st.session_state.token)
        if full_jd and len(full_jd) > len(prefill_jd):
            st.session_state[jd_key] = full_jd
            return full_jd, ""
        # Scraping failed — use what we have; only warn if we have nothing at all
        st.session_state[jd_key] = prefill_jd
        if not prefill_jd.strip():
            return prefill_jd, err or "Could not fetch the JD. Please paste it manually."
        return prefill_jd, ""

    st.session_state[jd_key] = prefill_jd
    return prefill_jd, ""


def _extract_file_text(uploaded_file) -> str:
    """Extract text from an uploaded PDF, DOCX, or TXT file."""
    name = uploaded_file.name.lower()
    raw = uploaded_file.read()
    if name.endswith(".txt") or name.endswith(".md"):
        return raw.decode("utf-8", errors="replace")
    if name.endswith(".docx"):
        try:
            import io
            from docx import Document
            doc = Document(io.BytesIO(raw))
            text = "\n".join(p.text for p in doc.paragraphs)
            if text.strip():
                return text
            st.warning("DOCX file appears empty. Please paste your resume instead.")
            return ""
        except ImportError:
            st.error("python-docx is not installed. Please paste your resume text instead.")
            return ""
        except Exception as e:
            st.error(f"Could not read DOCX file: {e}")
            return ""
    if name.endswith(".pdf"):
        try:
            import io
            from PyPDF2 import PdfReader
            reader = PdfReader(io.BytesIO(raw))
            text = "\n".join(p.extract_text() or "" for p in reader.pages)
            if text.strip():
                return text
            st.warning("PDF file appears empty or is scanned. Please paste your resume instead.")
            return ""
        except ImportError:
            st.error("PyPDF2 is not installed. Please paste your resume text instead.")
            return ""
        except Exception as e:
            st.error(f"Could not read PDF file: {e}")
            return ""
    return raw.decode("utf-8", errors="replace")


def resume_input(label: str = "Your Resume", prefill: str = "", key_prefix: str = "r", height: int = 350) -> str:
    """Shared resume input: file upload + text area + save option.

    Returns the resume text (from file, stored resume, or typed).
    """
    text_key = f"{key_prefix}_text"

    # File upload
    uploaded = st.file_uploader(
        "Upload resume (PDF, DOCX, or TXT)",
        type=["pdf", "docx", "txt", "md"],
        key=f"{key_prefix}_file_upload",
    )

    # When a new file is uploaded, extract → store in session + auto-push to DB if authed
    if uploaded:
        file_sig = f"{uploaded.name}_{uploaded.size}"
        if st.session_state.get(f"{key_prefix}_file_sig") != file_sig:
            extracted = _extract_file_text(uploaded)
            if extracted and not extracted.startswith("["):
                st.session_state[text_key] = extracted
                st.session_state[f"{key_prefix}_file_sig"] = file_sig
                st.session_state.stored_resume = extracted
                # Auto-push to DB for authenticated users — no button needed
                if is_authenticated():
                    if _upload_resume_to_cloud(st.session_state.token, extracted, uploaded.name):
                        st.session_state.resume_on_file = True
                st.rerun()

    # Pre-fill from stored resume or prefill (only if text_area hasn't been touched yet)
    if text_key not in st.session_state:
        default = prefill or st.session_state.get("stored_resume", "") or ""
        if default:
            st.session_state[text_key] = default

    # Text area — keyed, so Streamlit reads/writes st.session_state[text_key]
    resume_text = st.text_area(
        label,
        height=height,
        placeholder="Paste your resume text here, or upload a file above.",
        key=text_key,
    )

    # Status indicator + save/clear controls
    has_stored = bool(st.session_state.get("stored_resume", ""))
    if resume_text and resume_text.strip():
        if st.session_state.get("resume_on_file") and resume_text == st.session_state.stored_resume:
            st.markdown(
                '<span style="color: #22c55e; font-size: 12px;">&#10003; Resume saved to your account</span>',
                unsafe_allow_html=True,
            )
            col_clear, _ = st.columns([1, 3])
            with col_clear:
                if st.button("Remove saved resume", key=f"{key_prefix}_clear", help="Delete from account and clear field"):
                    st.session_state.stored_resume = ""
                    st.session_state.resume_on_file = False
                    st.session_state.pop(text_key, None)
                    if is_authenticated():
                        _delete_resume_from_cloud(st.session_state.token)
                    st.rerun()
        elif not has_stored:
            btn_label = "Save resume to account" if is_authenticated() else "Save as my resume"
            if st.button(btn_label, key=f"{key_prefix}_save", help="Auto-fills your resume on all pages"):
                st.session_state.stored_resume = resume_text
                if is_authenticated():
                    if _upload_resume_to_cloud(st.session_state.token, resume_text):
                        st.session_state.resume_on_file = True
                        st.success("Resume saved to your account.")
                    else:
                        st.warning("Saved for this session only (cloud sync failed).")
                else:
                    st.success("Resume saved for this session.")
                st.rerun()
        elif resume_text != st.session_state.stored_resume:
            col_a, col_b = st.columns(2)
            with col_a:
                btn_label = "Update saved resume" if not is_authenticated() else "Update & sync to account"
                if st.button(btn_label, key=f"{key_prefix}_update"):
                    st.session_state.stored_resume = resume_text
                    if is_authenticated():
                        if _upload_resume_to_cloud(st.session_state.token, resume_text):
                            st.session_state.resume_on_file = True
                    st.rerun()
            with col_b:
                if st.button("Clear saved resume", key=f"{key_prefix}_clear"):
                    st.session_state.stored_resume = ""
                    st.session_state.resume_on_file = False
                    st.session_state.pop(text_key, None)
                    if is_authenticated():
                        _delete_resume_from_cloud(st.session_state.token)
                    st.rerun()
        else:
            badge = "&#10003; Resume saved to your account" if st.session_state.get("resume_on_file") else "&#10003; Using saved resume"
            st.markdown(
                f'<span style="color: #22c55e; font-size: 12px;">{badge}</span>',
                unsafe_allow_html=True,
            )

    return resume_text


# ─── Navigation ──────────────────────────────────────────────────────────────

def render_nav():
    """Top navigation bar."""
    cols = st.columns([1, 1, 1, 1, 1, 1, 1, 2])

    with cols[0]:
        if st.button("Home", use_container_width=True, type="primary" if st.session_state.page == "home" else "secondary"):
            st.session_state.page = "home"
            st.rerun()
    with cols[1]:
        if st.button("Score", use_container_width=True, type="primary" if st.session_state.page == "scorer" else "secondary"):
            st.session_state.page = "scorer"
            st.rerun()
    with cols[2]:
        if st.button("Discover", use_container_width=True, type="primary" if st.session_state.page == "discover" else "secondary"):
            st.session_state.page = "discover"
            st.rerun()
    with cols[3]:
        if st.button("Rewrite", use_container_width=True, type="primary" if st.session_state.page == "rewriter" else "secondary"):
            st.session_state.page = "rewriter"
            st.rerun()
    with cols[4]:
        if st.button("Cover Letter", use_container_width=True, type="primary" if st.session_state.page == "cover_letter" else "secondary"):
            st.session_state.page = "cover_letter"
            st.rerun()
    with cols[5]:
        if is_authenticated():
            if st.button("Tracker", use_container_width=True, type="primary" if st.session_state.page == "tracker" else "secondary"):
                st.session_state.page = "tracker"
                st.rerun()
        else:
            st.button("Tracker", use_container_width=True, disabled=True, help="Sign in to use the Job Tracker")
    with cols[6]:
        if is_authenticated():
            if st.button("Dashboard", use_container_width=True, type="primary" if st.session_state.page == "dashboard" else "secondary"):
                st.session_state.page = "dashboard"
                st.rerun()
        else:
            if st.button("Login", use_container_width=True, type="primary" if st.session_state.page == "login" else "secondary"):
                st.session_state.page = "login"
                st.rerun()
    with cols[7]:
        if is_authenticated():
            if st.button("Logout", use_container_width=True):
                st.query_params.clear()
                st.session_state.token = None
                st.session_state.user = None
                st.session_state.stored_resume = ""
                st.session_state.resume_on_file = False
                st.session_state.page = "home"
                st.rerun()
        else:
            if st.button("Register", use_container_width=True, type="primary" if st.session_state.page == "register" else "secondary"):
                st.session_state.page = "register"
                st.rerun()

    st.markdown("---")


# ─── Plotly gauge charts ─────────────────────────────────────────────────────

def make_gauge(value: float, title: str) -> go.Figure:
    """Create a circular gauge chart for a score."""
    if value >= 75:
        bar_color = "#22c55e"
    elif value >= 50:
        bar_color = "#eab308"
    else:
        bar_color = "#ef4444"

    fig = go.Figure(go.Indicator(
        mode="gauge+number",
        value=value,
        number={"suffix": "%", "font": {"size": 48, "color": "#e2e8f0"}},
        title={"text": title, "font": {"size": 16, "color": "#94a3b8"}},
        gauge={
            "axis": {"range": [0, 100], "tickwidth": 0, "tickcolor": "#1e293b",
                     "tickfont": {"color": "#64748b"}},
            "bar": {"color": bar_color, "thickness": 0.3},
            "bgcolor": "#334155",
            "borderwidth": 0,
            "steps": [
                {"range": [0, 35], "color": "rgba(239,68,68,0.1)"},
                {"range": [35, 50], "color": "rgba(234,179,8,0.05)"},
                {"range": [50, 75], "color": "rgba(234,179,8,0.1)"},
                {"range": [75, 100], "color": "rgba(34,197,94,0.1)"},
            ],
            "threshold": {
                "line": {"color": "#6366f1", "width": 3},
                "thickness": 0.8,
                "value": value,
            },
        },
    ))
    fig.update_layout(
        height=250,
        margin={"t": 50, "b": 10, "l": 30, "r": 30},
        paper_bgcolor="rgba(0,0,0,0)",
        plot_bgcolor="rgba(0,0,0,0)",
        font={"color": "#e2e8f0"},
    )
    return fig


def make_bar_chart(labels: list, values: list, title: str) -> go.Figure:
    """Horizontal bar chart for score component breakdown."""
    colors = ["#22c55e" if v >= 70 else "#eab308" if v >= 40 else "#ef4444" for v in values]

    fig = go.Figure(go.Bar(
        x=values,
        y=labels,
        orientation="h",
        marker_color=colors,
        text=[f"{v:.0f}%" for v in values],
        textposition="auto",
        textfont={"color": "#e2e8f0", "size": 12},
    ))
    fig.update_layout(
        title={"text": title, "font": {"size": 16, "color": "#94a3b8"}, "x": 0},
        height=max(200, len(labels) * 40 + 80),
        margin={"t": 40, "b": 20, "l": 150, "r": 20},
        paper_bgcolor="rgba(0,0,0,0)",
        plot_bgcolor="rgba(0,0,0,0)",
        xaxis={"range": [0, 100], "showgrid": True, "gridcolor": "#1e293b",
               "tickfont": {"color": "#64748b"}, "title": ""},
        yaxis={"tickfont": {"color": "#e2e8f0"}, "autorange": "reversed"},
        bargap=0.3,
    )
    return fig


# ─── Score display helpers ───────────────────────────────────────────────────

def render_keyword_chips(matched: list, missing: list):
    """Render matched (green) and missing (red) keyword chips."""
    html = ""
    if matched:
        html += "<div style='margin-bottom: 8px;'><span style='color: #94a3b8; font-size: 13px;'>Matched:</span><br>"
        for kw in matched[:20]:
            html += f'<span class="chip-matched">{kw}</span>'
        html += "</div>"
    if missing:
        html += "<div><span style='color: #94a3b8; font-size: 13px;'>Missing:</span><br>"
        for kw in missing[:15]:
            html += f'<span class="chip-missing">{kw}</span>'
        html += "</div>"
    if html:
        st.markdown(html, unsafe_allow_html=True)


def render_hr_badge(recommendation: str):
    """Render colored HR recommendation badge."""
    rec_upper = recommendation.upper() if recommendation else ""
    if "STRONG" in rec_upper or "INTERVIEW" in rec_upper:
        css_class = "badge-green"
    elif "MAYBE" in rec_upper:
        css_class = "badge-yellow"
    else:
        css_class = "badge-red"

    st.markdown(
        f'<span class="badge {css_class}">{recommendation}</span>',
        unsafe_allow_html=True,
    )


def render_quick_wins(explanation: dict):
    """Render improvement suggestions from ATS explanation."""
    quick_wins = explanation.get("quick_wins", [])
    missing_kw = explanation.get("top_missing_keywords", [])

    if quick_wins:
        st.markdown("##### Quick Wins")
        for win in quick_wins:
            st.markdown(f"- {win}")

    if missing_kw:
        st.markdown("##### Top Missing Keywords")
        for item in missing_kw[:5]:
            delta = item.get("estimated_score_increase", "")
            placement = item.get("suggested_placement", "")
            st.markdown(f"- **{item['keyword']}** — add to _{placement}_ ({delta})")


def render_hr_insights(explanation: dict):
    """Render HR improvement suggestions."""
    improvements = explanation.get("priority_improvements", [])
    strengths = explanation.get("strengths_to_emphasize", [])
    risks = explanation.get("risk_mitigations", [])

    if strengths:
        st.markdown("##### Your Strengths")
        for s in strengths:
            st.markdown(f"- **{s['factor']}** ({s['current_score']:.0f}%) — {s['advice']}")

    if improvements:
        st.markdown("##### Priority Improvements")
        for imp in improvements:
            st.markdown(f"- **{imp['factor']}** ({imp['current_score']:.0f}%) — {imp['suggestion']}")

    if risks:
        st.markdown("##### Risk Mitigations")
        for r in risks:
            st.markdown(f"- **{r['risk']}** ({r['penalty']}) — {r['mitigation']}")


# ─── Pages ───────────────────────────────────────────────────────────────────

def page_home():
    """Landing / hero page."""
    # Hero
    st.markdown("""
    <div style="text-align: center; padding: 40px 0 20px 0;">
        <h1 style="color: #e2e8f0; font-size: 42px; font-weight: 800; margin-bottom: 8px;">
            Score Your Resume in Seconds
        </h1>
        <p style="color: #94a3b8; font-size: 18px; max-width: 600px; margin: 0 auto;">
            AI-powered ATS + HR scoring that tells you exactly how your resume will perform
            with applicant tracking systems and human recruiters.
        </p>
    </div>
    """, unsafe_allow_html=True)

    # Benefit cards
    col1, col2, col3 = st.columns(3)
    with col1:
        st.markdown("""
        <div class="stat-card">
            <div class="stat-num">ATS</div>
            <div class="stat-label">Keyword & Semantic Analysis</div>
        </div>
        """, unsafe_allow_html=True)
        st.markdown("""
        <div class="card" style="min-height: 120px;">
            <p style="color: #e2e8f0; font-weight: 600; margin-bottom: 8px;">Beat the Bots</p>
            <p style="color: #94a3b8; font-size: 14px;">
                8-component scoring: keywords, phrases, semantic matching, BM25 ranking,
                industry terms, skill graphs, job title match, and recency weighting.
            </p>
        </div>
        """, unsafe_allow_html=True)
    with col2:
        st.markdown("""
        <div class="stat-card">
            <div class="stat-num">HR</div>
            <div class="stat-label">Recruiter Simulation</div>
        </div>
        """, unsafe_allow_html=True)
        st.markdown("""
        <div class="card" style="min-height: 120px;">
            <p style="color: #e2e8f0; font-weight: 600; margin-bottom: 8px;">Impress Humans</p>
            <p style="color: #94a3b8; font-size: 14px;">
                6-factor recruiter evaluation: experience fit, skills match, career trajectory,
                impact signals, competitive edge, and job fit analysis.
            </p>
        </div>
        """, unsafe_allow_html=True)
    with col3:
        st.markdown("""
        <div class="stat-card">
            <div class="stat-num">5</div>
            <div class="stat-label">Free Scores</div>
        </div>
        """, unsafe_allow_html=True)
        st.markdown("""
        <div class="card" style="min-height: 120px;">
            <p style="color: #e2e8f0; font-weight: 600; margin-bottom: 8px;">Start Free</p>
            <p style="color: #94a3b8; font-size: 14px;">
                No account needed for your first 5 scores. Upgrade to Pro ($12/month)
                for unlimited scoring with detailed explanations.
            </p>
        </div>
        """, unsafe_allow_html=True)

    st.markdown("<br>", unsafe_allow_html=True)

    # CTA
    _, cta_col, _ = st.columns([2, 3, 2])
    with cta_col:
        if st.button("Score My Resume Now", use_container_width=True, type="primary"):
            st.session_state.page = "scorer"
            st.rerun()

    st.markdown("<br>", unsafe_allow_html=True)

    # Pricing section
    render_pricing()


def render_pricing():
    """Pricing comparison section — Free / Pro / Ultra."""
    st.markdown('<div class="section-title" style="text-align: center;">Pricing</div>', unsafe_allow_html=True)
    st.markdown('<div class="section-subtitle" style="text-align: center;">Simple, transparent pricing</div>', unsafe_allow_html=True)

    col1, col2, col3 = st.columns(3)
    with col1:
        st.markdown("""
        <div class="card" style="text-align: center; min-height: 380px;">
            <p style="color: #94a3b8; font-size: 14px; text-transform: uppercase; letter-spacing: 1px;">Free</p>
            <div class="price-tag">$0</div>
            <hr style="border-color: #334155; margin: 16px 0;">
            <div style="text-align: left; color: #94a3b8; font-size: 14px; line-height: 2;">
                &#10003;&nbsp; 5 total scores<br>
                &#10003;&nbsp; ATS + HR scoring<br>
                &#10003;&nbsp; Domain auto-detection<br>
                &#10007;&nbsp; LLM-augmented scoring<br>
                &#10007;&nbsp; AI resume rewriting
            </div>
        </div>
        """, unsafe_allow_html=True)

    with col2:
        st.markdown("""
        <div class="card-accent" style="text-align: center; min-height: 380px;">
            <p style="color: #818cf8; font-size: 14px; text-transform: uppercase; letter-spacing: 1px;">Pro</p>
            <div class="price-tag">$12<span class="price-period">/month</span></div>
            <hr style="border-color: #6366f1; margin: 16px 0;">
            <div style="text-align: left; color: #94a3b8; font-size: 14px; line-height: 2;">
                &#10003;&nbsp; Unlimited scores<br>
                &#10003;&nbsp; ATS + HR + LLM scoring<br>
                &#10003;&nbsp; Detailed explanations<br>
                &#10003;&nbsp; API key access<br>
                &#10003;&nbsp; 10 AI rewrites/month
            </div>
        </div>
        """, unsafe_allow_html=True)

    with col3:
        st.markdown("""
        <div class="card-accent" style="text-align: center; min-height: 380px; border-color: #a855f7;">
            <p style="color: #a855f7; font-size: 14px; text-transform: uppercase; letter-spacing: 1px;">Ultra</p>
            <div class="price-tag" style="color: #c084fc;">$29<span class="price-period">/month</span></div>
            <hr style="border-color: #a855f7; margin: 16px 0;">
            <div style="text-align: left; color: #94a3b8; font-size: 14px; line-height: 2;">
                &#10003;&nbsp; Everything in Pro<br>
                &#10003;&nbsp; Unlimited AI rewrites<br>
                &#10003;&nbsp; Before/after scoring<br>
                &#10003;&nbsp; Download tailored DOCX<br>
                &#10003;&nbsp; Change tracking
            </div>
        </div>
        """, unsafe_allow_html=True)


def page_scorer():
    """Interactive resume scorer — paste resume + JD, get visual results."""
    st.markdown('<div class="section-title">Score Your Resume</div>', unsafe_allow_html=True)
    st.markdown(
        '<div class="section-subtitle">Paste your resume and job description to see how well they match.</div>',
        unsafe_allow_html=True,
    )

    # Free scores counter / Pro LLM toggle
    is_paid = is_authenticated() and st.session_state.user and st.session_state.user.get("tier") in ("pro", "ultra")

    if is_paid:
        use_llm = st.toggle("Include LLM Analysis (Pro)", value=True, help="Add Claude-powered AI analysis on top of ATS + HR scoring")
    else:
        use_llm = False
        remaining = max(0, 5 - st.session_state.scores_used)
        st.markdown(
            f'<div style="text-align: right; color: #94a3b8; font-size: 13px; margin-bottom: 12px;">'
            f'Free scores remaining: <strong style="color: #818cf8;">{remaining}/5</strong></div>',
            unsafe_allow_html=True,
        )
        # Signup prompt for anonymous users
        if not is_authenticated():
            signup_col1, signup_col2 = st.columns([5, 2])
            with signup_col1:
                st.markdown(
                    '<p style="color: #94a3b8; font-size: 14px; margin: 0;">'
                    'Create a free account to track your scores and unlock Pro features.</p>',
                    unsafe_allow_html=True,
                )
            with signup_col2:
                if st.button("Sign Up Free", key="top_signup", type="secondary", use_container_width=True):
                    st.session_state.page = "register"
                    st.rerun()

    # Pre-fill from Discover page handoff (if any)
    prefill_resume = st.session_state.pop("prefill_resume", "") or ""
    prefill_jd = st.session_state.pop("prefill_jd", "") or ""
    prefill_job_url = st.session_state.pop("prefill_job_url", "") or ""
    prefill_job_title = st.session_state.pop("prefill_job_title", "") or ""
    if prefill_jd or prefill_job_url:
        prefill_jd, jd_warn = _apply_jd_prefill("score_jd_text", prefill_jd, prefill_job_url, prefill_job_title)
        if jd_warn:
            st.warning(f"⚠ {jd_warn}")

    # Input columns
    col_resume, col_jd = st.columns(2)
    with col_resume:
        resume_text = resume_input("Resume", prefill=prefill_resume, key_prefix="score")
    with col_jd:
        jd_text = st.text_area(
            "Job Description",
            height=350,
            placeholder="Paste the job description here...\n\nInclude requirements, responsibilities, and qualifications.",
            key="score_jd_text",
        )

    # Score button
    _, btn_col, _ = st.columns([3, 2, 3])
    with btn_col:
        score_clicked = st.button("Analyze Resume", use_container_width=True, type="primary")

    if score_clicked:
        if not resume_text or not jd_text:
            st.error("Please paste both your resume and the job description.")
            return
        if len(resume_text.strip()) < 100:
            st.warning("Resume seems too short. Please paste the full text.")
            return
        if len(jd_text.strip()) < 50:
            st.warning("Job description seems too short. Please paste the full listing.")
            return

        endpoint = "/score/combined" if use_llm else "/score/both"
        payload = {"resume_text": resume_text, "jd_text": jd_text}
        if not use_llm:
            payload["include_explanation"] = True

        _stage_labels = {
            "scoring":     "Running ATS + HR analysis in parallel…",
            "llm_scoring": "Claude AI is reviewing your resume…",
        }
        _status_box = st.empty()
        _pbar = st.progress(0)
        _score_data = None
        _score_error = None
        _score_status = 200

        for _ev in api_stream(endpoint, payload, token=st.session_state.token):
            _stage = _ev.get("stage", "")
            _pct   = _ev.get("pct", 0)

            if _stage == "error":
                _score_error  = _ev.get("detail", "Unknown error")
                _score_status = _ev.get("status", 500)
                break
            elif _stage == "done":
                _score_data = _ev.get("result")
                break
            else:
                _label = _stage_labels.get(_stage, "Analyzing…")
                _status_box.markdown(
                    f'<div style="background:#1e293b;border:1px solid #334155;'
                    f'border-radius:8px;padding:14px 18px;color:#94a3b8;font-size:14px;">'
                    f'<span style="color:#818cf8;font-weight:600;">Scorer</span>'
                    f'&nbsp;·&nbsp;{_label}</div>',
                    unsafe_allow_html=True,
                )
                _pbar.progress(max(1, _pct) / 100)

        _status_box.empty()
        _pbar.empty()

        if _score_status == 402:
            st.error("You've used all 5 free scores.")
            st.markdown("---")
            st.markdown(
                "**Upgrade to Pro** for unlimited scoring, detailed explanations, and LLM-augmented analysis."
            )
            if is_authenticated():
                if st.button("Upgrade to Pro ($12/month)", type="primary", use_container_width=True):
                    checkout = api("POST", "/billing/checkout", token=st.session_state.token)
                    if checkout["status"] == 200 and "checkout_url" in checkout["data"]:
                        st.link_button("Complete Payment on Stripe", checkout["data"]["checkout_url"], use_container_width=True)
                    else:
                        st.error(checkout["data"].get("detail", "Could not create checkout session."))
            else:
                st.info("Create a free account first, then upgrade to Pro.")
                if st.button("Create Account", type="primary"):
                    st.session_state.page = "register"
                    st.rerun()
            return

        if _score_error:
            st.error(f"Scoring failed: {_score_error}")
            return

        if not _score_data:
            st.error("Scorer returned no data. Please try again.")
            return

        result = {"status": 200, "data": _score_data}

        # Success — store result and bump counter
        data = result["data"]
        # Normalize /score/combined response to match /score/both shape
        if use_llm and "rules_ats" in data:
            data["_llm_mode"] = True
            data["ats"] = data.get("rules_ats", {})
            data["hr"] = data.get("rules_hr", {}) or {}
            data["summary"] = {
                "ats_score": data.get("combined_ats", data["ats"].get("total_score", 0)),
                "hr_score": data.get("combined_hr", data["hr"].get("overall_score", 0)),
                "ats_rating": data["ats"].get("rating", ""),
                "hr_recommendation": data["hr"].get("recommendation", ""),
                "overall_assessment": "",
            }
        st.session_state.score_result = data
        st.session_state.scores_used += 1

    # Display results if available
    data = st.session_state.score_result
    if data:
        render_score_results(data)

        # Post-score signup/upgrade prompt
        if not is_authenticated():
            st.markdown("---")
            st.markdown("""
            <div class="card-accent" style="text-align: center;">
                <p style="color: #818cf8; font-weight: 700; font-size: 18px; margin-bottom: 4px;">
                    Want unlimited scores + AI-powered analysis?
                </p>
                <p style="color: #94a3b8; font-size: 14px;">
                    Create a free account to save your results, then upgrade to Pro ($12/month)
                    for unlimited scoring and Claude-powered LLM analysis.
                </p>
            </div>
            """, unsafe_allow_html=True)
            _, cta_col, _ = st.columns([2, 3, 2])
            with cta_col:
                if st.button("Create Free Account", key="post_score_signup", type="primary", use_container_width=True):
                    st.session_state.page = "register"
                    st.rerun()
        elif not is_pro:
            st.markdown("---")
            st.markdown("""
            <div class="card-accent" style="text-align: center;">
                <p style="color: #818cf8; font-weight: 700; font-size: 18px; margin-bottom: 4px;">
                    Unlock LLM-Powered Analysis
                </p>
                <p style="color: #94a3b8; font-size: 14px;">
                    Upgrade to Pro ($12/month) for unlimited scoring and Claude AI analysis
                    with detailed evidence and improvement suggestions.
                </p>
            </div>
            """, unsafe_allow_html=True)
            _, cta_col, _ = st.columns([2, 3, 2])
            with cta_col:
                if st.button("Upgrade to Pro", key="post_score_upgrade", type="primary", use_container_width=True):
                    st.session_state.page = "dashboard"
                    st.rerun()


def render_score_results(data: dict):
    """Render visual score breakdown from API response."""
    ats = data.get("ats", {})
    hr = data.get("hr", {})
    summary = data.get("summary", {})
    explanation = data.get("explanation", {})

    ats_score = summary.get("ats_score", ats.get("total_score", 0))
    hr_score = summary.get("hr_score", hr.get("overall_score", 0))

    st.markdown("---")

    # ─── Gauges ──────────────────────────────────────────────────────────
    gauge_col1, gauge_col2 = st.columns(2)
    with gauge_col1:
        st.plotly_chart(make_gauge(ats_score, "ATS Score"), use_container_width=True)
    with gauge_col2:
        st.plotly_chart(make_gauge(hr_score, "HR Score"), use_container_width=True)

    # Overall assessment
    assessment = summary.get("overall_assessment", "")
    if assessment:
        st.markdown(
            f'<div style="text-align: center; color: #94a3b8; font-size: 16px; margin-bottom: 16px;">'
            f'{assessment}</div>',
            unsafe_allow_html=True,
        )

    # HR recommendation badge
    hr_rec = summary.get("hr_recommendation", hr.get("recommendation", ""))
    if hr_rec:
        st.markdown(
            '<div style="text-align: center; margin-bottom: 24px;">',
            unsafe_allow_html=True,
        )
        render_hr_badge(hr_rec)
        st.markdown('</div>', unsafe_allow_html=True)

    # ─── Detailed Tabs ───────────────────────────────────────────────────
    llm_data = data.get("llm") if data.get("_llm_mode") else None
    if llm_data and not llm_data.get("error"):
        tab_ats, tab_hr, tab_llm = st.tabs(["ATS Analysis", "HR Analysis", "LLM Analysis"])
    else:
        tab_ats, tab_hr = st.tabs(["ATS Analysis", "HR Analysis"])
        tab_llm = None

    with tab_ats:
        render_ats_tab(ats, explanation.get("ats", {}))

    with tab_hr:
        render_hr_tab(hr, explanation.get("hr", {}))

    if tab_llm is not None:
        with tab_llm:
            render_llm_tab(llm_data, data.get("blend_details", {}))


def render_ats_tab(ats: dict, ats_explanation: dict):
    """ATS detailed breakdown tab."""
    # Component bar chart
    component_map = {
        "Keyword Match": ats.get("keyword_score", 0),
        "Phrase Match": ats.get("phrase_score", 0),
        "Industry Terms": ats.get("weighted_score", 0),
        "Semantic Similarity": ats.get("semantic_score", 0),
        "BM25 Relevance": ats.get("bm25_score", 0),
        "Job Title Match": ats.get("job_title_score", 0),
        "Graph Centrality": ats.get("graph_score", 0),
        "Skill Recency": ats.get("recency_score", 0),
    }
    # Filter out zero/missing components
    labels = [k for k, v in component_map.items() if v is not None]
    values = [component_map[k] for k in labels]

    if labels:
        st.plotly_chart(
            make_bar_chart(labels, values, "ATS Component Breakdown"),
            use_container_width=True,
        )

    # Keyword chips
    matched = ats.get("matched_keywords", [])
    missing = ats.get("missing_keywords", [])
    if matched or missing:
        st.markdown("##### Keywords")
        render_keyword_chips(matched, missing)

    # Matched phrases
    matched_phrases = ats.get("matched_phrases", [])
    if matched_phrases:
        st.markdown("##### Matched Phrases")
        phrase_html = "".join(f'<span class="chip-matched">{p}</span>' for p in matched_phrases[:15])
        st.markdown(phrase_html, unsafe_allow_html=True)

    # Domain detection
    domain = ats.get("domain", "")
    if domain and isinstance(domain, str):
        st.markdown(f"**Detected Domain:** {domain.replace('_', ' ').title()}")

    # Readability
    readability = ats.get("readability", {})
    if isinstance(readability, dict) and readability.get("flesch_kincaid_grade"):
        grade = readability["flesch_kincaid_grade"]
        st.markdown(f"**Readability:** Grade {grade:.1f} (optimal: 10-12)")

    # Format risk
    format_risk = ats.get("format_risk_score")
    if format_risk is not None:
        risk_label = "Low" if format_risk < 20 else "Medium" if format_risk < 50 else "High"
        st.markdown(f"**Format Risk:** {format_risk:.0f}% ({risk_label})")

    # ATS rating
    rating = ats.get("rating", "")
    likelihood = ats.get("likelihood", "")
    if rating:
        st.markdown(f"**Rating:** {rating} — {likelihood}")

    # Quick wins from explanation
    if ats_explanation:
        st.markdown("---")
        render_quick_wins(ats_explanation)


def render_hr_tab(hr: dict, hr_explanation: dict):
    """HR detailed breakdown tab."""
    # Factor bar chart
    breakdown = hr.get("factor_breakdown", {})
    factor_labels = {
        "experience": "Experience Fit",
        "skills": "Skills Match",
        "trajectory": "Career Trajectory",
        "impact": "Impact Signals",
        "competitive": "Competitive Edge",
        "job_fit": "Job Fit",
    }

    labels = []
    values = []
    for key, label in factor_labels.items():
        val = breakdown.get(key)
        if val is not None and isinstance(val, (int, float)):
            labels.append(label)
            values.append(val)

    if labels:
        st.plotly_chart(
            make_bar_chart(labels, values, "HR Factor Breakdown"),
            use_container_width=True,
        )

    # Penalties
    penalties = hr.get("penalties_applied", {})
    if penalties:
        active_penalties = {k: v for k, v in penalties.items() if v and v > 0}
        if active_penalties:
            st.markdown("##### Penalties Applied")
            for name, val in active_penalties.items():
                st.markdown(f"- **{name.replace('_', ' ').title()}**: -{val:.1f} points")

    # Visual / F-pattern score
    visual = hr.get("visual_score")
    if visual is not None:
        st.markdown(f"**Visual/F-Pattern Score:** {visual:+.1f} points")

    # Metrics density
    metrics_density = hr.get("metrics_density")
    if metrics_density is not None:
        pct = metrics_density * 100 if metrics_density <= 1 else metrics_density
        st.markdown(f"**Metrics Density:** {pct:.0f}% of bullets have quantified impact")

    # HR explanation insights
    if hr_explanation:
        st.markdown("---")
        render_hr_insights(hr_explanation)


def render_llm_tab(llm_data: dict, blend_details: dict):
    """LLM (Claude) analysis tab — Pro only."""
    llm_ats = llm_data.get("ats_score")
    llm_hr = llm_data.get("hr_score")

    # LLM scores gauges
    if llm_ats is not None and llm_hr is not None:
        g1, g2 = st.columns(2)
        with g1:
            st.plotly_chart(make_gauge(llm_ats, "LLM ATS Score"), use_container_width=True)
        with g2:
            st.plotly_chart(make_gauge(llm_hr, "LLM HR Score"), use_container_width=True)

    # Blend details
    if blend_details and blend_details.get("method") != "rules_only":
        st.markdown(
            '<p style="color: #94a3b8; font-size: 13px; text-align: center;">'
            "Scores above are from Claude AI. The combined scores on the gauges blend rules-based and LLM analysis.</p>",
            unsafe_allow_html=True,
        )

    # Explanation summary
    explanation = llm_data.get("explanation", "")
    if explanation:
        st.markdown("##### AI Summary")
        st.info(explanation)

    # Dimension breakdown — ATS
    dimensions = llm_data.get("dimensions", {})
    ats_dims = dimensions.get("ats", {})
    if ats_dims:
        dim_labels_ats = {
            "keyword_match": "Keyword Match",
            "phrase_match": "Phrase Match",
            "industry_terms": "Industry Terms",
            "semantic_similarity": "Semantic Similarity",
            "bm25_relevance": "BM25 Relevance",
            "graph_centrality": "Graph Centrality",
            "skill_recency": "Skill Recency",
            "job_title_match": "Job Title Match",
        }
        labels = []
        values = []
        for key, label in dim_labels_ats.items():
            dim = ats_dims.get(key, {})
            if isinstance(dim, dict) and "score" in dim:
                labels.append(label)
                values.append(dim["score"] * 20)  # 0-5 scale → 0-100
        if labels:
            st.plotly_chart(make_bar_chart(labels, values, "LLM ATS Dimensions"), use_container_width=True)

        # Evidence details
        with st.expander("ATS Dimension Evidence"):
            for key, label in dim_labels_ats.items():
                dim = ats_dims.get(key, {})
                if isinstance(dim, dict) and dim.get("evidence"):
                    st.markdown(f"**{label}** ({dim.get('score', 0)}/5): {dim['evidence']}")

    # Dimension breakdown — HR
    hr_dims = dimensions.get("hr", {})
    if hr_dims:
        dim_labels_hr = {
            "job_fit": "Job Fit",
            "experience_fit": "Experience Fit",
            "skills_in_action": "Skills in Action",
            "impact_signals": "Impact Signals",
            "career_trajectory": "Career Trajectory",
            "competitive_edge": "Competitive Edge",
        }
        labels = []
        values = []
        for key, label in dim_labels_hr.items():
            dim = hr_dims.get(key, {})
            if isinstance(dim, dict) and "score" in dim:
                labels.append(label)
                values.append(dim["score"] * 20)
        if labels:
            st.plotly_chart(make_bar_chart(labels, values, "LLM HR Dimensions"), use_container_width=True)

        with st.expander("HR Dimension Evidence"):
            for key, label in dim_labels_hr.items():
                dim = hr_dims.get(key, {})
                if isinstance(dim, dict) and dim.get("evidence"):
                    st.markdown(f"**{label}** ({dim.get('score', 0)}/5): {dim['evidence']}")

    # Penalties
    penalties = llm_data.get("hr_penalties", {})
    if penalties:
        hopping = penalties.get("job_hopping", 0)
        gaps = penalties.get("gaps", 0)
        notes = penalties.get("notes", "")
        if hopping or gaps:
            st.markdown("##### LLM Risk Penalties")
            if hopping:
                st.markdown(f"- **Job Hopping**: {hopping} points")
            if gaps:
                st.markdown(f"- **Gaps**: {gaps} points")
            if notes:
                st.markdown(f"- _{notes}_")

    # Domain detected
    domain = llm_data.get("domain_detected", "")
    if domain:
        st.markdown(f"**Domain Detected (LLM):** {domain.replace('_', ' ').title()}")

    st.markdown(f"**Model:** {llm_data.get('model_used', 'claude-sonnet-4-6')}")


# ─── Rewriter page (Ultra) ──────────────────────────────────────────────────

def page_rewriter():
    """Pro + Ultra tier: AI resume rewriting with before/after scores."""
    st.markdown('<div class="section-title">AI Resume Rewriter</div>', unsafe_allow_html=True)
    st.markdown(
        '<div class="section-subtitle">Upload your resume and job description. '
        'Claude AI will tailor your resume and show before/after scores.</div>',
        unsafe_allow_html=True,
    )

    user_tier = ""
    if is_authenticated() and st.session_state.user:
        user_tier = st.session_state.user.get("tier", "free")

    can_rewrite = user_tier in ("pro", "ultra")

    if not can_rewrite:
        # Consume prefill session state so it doesn't carry over on next navigation
        st.session_state.pop("prefill_resume", None)
        st.session_state.pop("prefill_jd", None)
        st.session_state.pop("prefill_job_url", None)
        st.session_state.pop("prefill_job_title", None)
        # Show upgrade prompt for free/anonymous users
        st.markdown("""
        <div class="card-accent" style="text-align: center; border-color: #818cf8;">
            <p style="color: #818cf8; font-weight: 700; font-size: 20px; margin-bottom: 8px;">
                Pro Feature
            </p>
            <p style="color: #94a3b8; font-size: 15px; margin-bottom: 16px;">
                AI resume rewriting is available on Pro ($12/month, 10 rewrites) and
                Ultra ($29/month, unlimited). Claude AI tailors your resume to match
                the job description while keeping your real experience intact.
            </p>
            <p style="color: #94a3b8; font-size: 14px;">
                &#10003; Before/after ATS + HR scores &nbsp;&nbsp;
                &#10003; Download tailored resume &nbsp;&nbsp;
                &#10003; Change tracking
            </p>
        </div>
        """, unsafe_allow_html=True)

        st.markdown("")
        if not is_authenticated():
            _, cta_col, _ = st.columns([2, 3, 2])
            with cta_col:
                if st.button("Sign Up to Get Started", type="primary", use_container_width=True):
                    st.session_state.page = "register"
                    st.rerun()
        else:
            _, pro_col, ultra_col, _ = st.columns([1, 2, 2, 1])
            with pro_col:
                if st.button("Pro — $12/month", type="primary", use_container_width=True):
                    with st.spinner("Creating checkout session..."):
                        result = api("POST", "/billing/checkout", {"tier": "pro"}, token=st.session_state.token)
                    if result["status"] == 200 and "checkout_url" in result["data"]:
                        st.link_button("Complete Payment on Stripe", result["data"]["checkout_url"], use_container_width=True)
                    elif result["status"] == 503:
                        st.warning("Stripe billing is not configured yet.")
                    else:
                        st.error(result["data"].get("detail", "Could not create checkout session."))
            with ultra_col:
                if st.button("Ultra — $29/month", use_container_width=True):
                    with st.spinner("Creating checkout session..."):
                        result = api("POST", "/billing/checkout", {"tier": "ultra"}, token=st.session_state.token)
                    if result["status"] == 200 and "checkout_url" in result["data"]:
                        st.link_button("Complete Payment on Stripe", result["data"]["checkout_url"], use_container_width=True)
                    elif result["status"] == 503:
                        st.warning("Stripe billing is not configured yet.")
                    else:
                        st.error(result["data"].get("detail", "Could not create checkout session."))

        st.markdown("---")
        render_pricing()
        return

    # ─── Pro / Ultra user: show the rewriter ──────────────────────────────
    # Show remaining rewrites for Pro users
    if user_tier == "pro":
        usage = api("GET", "/auth/usage", token=st.session_state.token)
        if usage["status"] == 200:
            rewrites_info = usage["data"].get("rewrites", {})
            remaining = rewrites_info.get("remaining", 10)
            limit = rewrites_info.get("limit", 10)
            st.markdown(
                f'<div style="text-align: right; color: #94a3b8; font-size: 13px; margin-bottom: 12px;">'
                f'Rewrites remaining: <strong style="color: #818cf8;">{remaining}/{limit}</strong> this month</div>',
                unsafe_allow_html=True,
            )
    # Pre-fill from Discover page handoff (if any)
    prefill_resume = st.session_state.pop("prefill_resume", "") or ""
    prefill_jd = st.session_state.pop("prefill_jd", "") or ""
    prefill_job_url = st.session_state.pop("prefill_job_url", "") or ""
    prefill_job_title = st.session_state.pop("prefill_job_title", "") or ""
    prefill_company = st.session_state.pop("prefill_company", "") or ""
    if prefill_jd or prefill_job_url:
        prefill_jd, jd_warn = _apply_jd_prefill("rewrite_jd_text", prefill_jd, prefill_job_url, prefill_job_title)
        if jd_warn:
            st.warning(f"⚠ {jd_warn}")

    # Company + job title fields (used for file naming and tracker)
    _info_col1, _info_col2 = st.columns(2)
    with _info_col1:
        _company_input = st.text_input(
            "Company Name",
            value=st.session_state.get("rewrite_company_val", prefill_company),
            placeholder="e.g. Pfizer, Google, Mayo Clinic…",
            key="rewrite_company_input",
        )
    with _info_col2:
        _job_title_input = st.text_input(
            "Job Title",
            value=st.session_state.get("rewrite_job_title_val", prefill_job_title),
            placeholder="e.g. Data Scientist, Clinical Research Associate…",
            key="rewrite_job_title_input",
        )

    col_resume, col_jd = st.columns(2)
    with col_resume:
        resume_text = resume_input("Your Resume", prefill=prefill_resume, key_prefix="rewrite", height=400)
    with col_jd:
        jd_text = st.text_area(
            "Target Job Description",
            height=400,
            placeholder="Paste the job description you're applying for...",
            key="rewrite_jd_text",
        )

    # ── Resume format selector ─────────────────────────────────────────────
    st.markdown("#### Choose Resume Format")
    _FORMAT_OPTIONS = {
        "ats":       ("ATS Standard",      "Clean, keyword-optimised. Passes applicant tracking systems. Best for most corporate roles."),
        "harvard":   ("Harvard Classic",   "Traditional academic/professional. Times New Roman, underlined section headers. Best for law, academia, consulting."),
        "modern":    ("Modern Clean",      "Contemporary design with blue accent lines. Concise, impact-focused bullets. Best for tech, startups, marketing."),
        "executive": ("Executive Classic", "Formal, centered layout with Georgia serif font. Emphasis on leadership scale. Best for C-suite and senior finance."),
    }
    _fmt_cols = st.columns(4)
    _selected_fmt = st.session_state.get("rewrite_format", "ats")
    for _ci, (_fkey, (_fname, _fdesc)) in enumerate(zip(_FORMAT_OPTIONS, _FORMAT_OPTIONS.values())):
        with _fmt_cols[_ci]:
            _is_sel = (_fkey == _selected_fmt)
            _border = "#818cf8" if _is_sel else "#334155"
            _bg     = "#1e1b4b" if _is_sel else "#1e293b"
            st.markdown(
                f'<div style="border:2px solid {_border};border-radius:8px;padding:10px 12px;'
                f'background:{_bg};cursor:pointer;min-height:90px;">'
                f'<div style="color:#e2e8f0;font-weight:700;font-size:13px;margin-bottom:4px;">{_fname}</div>'
                f'<div style="color:#94a3b8;font-size:11px;line-height:1.4;">{_fdesc}</div>'
                f'</div>',
                unsafe_allow_html=True,
            )
            if st.button(
                "Select" if not _is_sel else "✓ Selected",
                key=f"fmt_{_fkey}",
                use_container_width=True,
                type="primary" if _is_sel else "secondary",
            ):
                st.session_state.rewrite_format = _fkey
                st.rerun()

    _selected_fmt = st.session_state.get("rewrite_format", "ats")
    _fmt_name = _FORMAT_OPTIONS[_selected_fmt][0]
    st.markdown(
        f'<div style="color:#818cf8;font-size:13px;margin-top:4px;margin-bottom:8px;">'
        f'Format: <strong>{_fmt_name}</strong> — Claude will structure your resume in this style '
        f'and the DOCX will use its fonts and layout.</div>',
        unsafe_allow_html=True,
    )

    _, btn_col, _ = st.columns([3, 2, 3])
    with btn_col:
        rewrite_clicked = st.button("Rewrite My Resume", use_container_width=True, type="primary")

    if rewrite_clicked:
        if not resume_text or not jd_text:
            st.error("Please paste both your resume and the job description.")
            return
        if len(resume_text.strip()) < 100:
            st.warning("Resume seems too short. Please paste the full text.")
            return

        # ── Stream SSE progress from /rewrite endpoint ────────────────────────
        _stage_labels = {
            "scoring_original":  "Scoring your original resume…",
            "rewriting":         f"Claude AI is rewriting in {_fmt_name} format…",
            "scoring_rewritten": "Scoring the tailored resume (ATS + HR)…",
            "scoring_llm":       "Running LLM evaluation…",
        }
        _status_box = st.empty()
        _pbar = st.progress(0)
        _rewrite_result = None
        _rewrite_error = None

        for _event in api_stream(
            "/rewrite",
            {"resume_text": resume_text, "jd_text": jd_text, "format_style": _selected_fmt},
            token=st.session_state.token,
        ):
            _stage = _event.get("stage", "")
            _pct   = _event.get("pct", 0)

            if _stage == "error":
                _rewrite_error = _event.get("detail", "Unknown error")
                break
            elif _stage == "done":
                _rewrite_result = _event.get("result")
                break
            else:
                _label = _stage_labels.get(_stage, "Processing…")
                _status_box.markdown(
                    f'<div style="background:#1e293b;border:1px solid #334155;'
                    f'border-radius:8px;padding:14px 18px;color:#94a3b8;font-size:14px;">'
                    f'<span style="color:#818cf8;font-weight:600;">Claude AI</span>'
                    f'&nbsp;·&nbsp;{_label}</div>',
                    unsafe_allow_html=True,
                )
                _pbar.progress(max(1, _pct) / 100)

        _status_box.empty()
        _pbar.empty()

        if _rewrite_error:
            st.error(f"Rewrite failed: {_rewrite_error}")
            return
        if not _rewrite_result:
            st.error("Rewrite returned no data. Please try again.")
            return

        st.session_state.rewrite_result = _rewrite_result
        st.session_state.rewrite_jd_saved = jd_text
        st.session_state.rewrite_company_val = _company_input
        st.session_state.rewrite_job_title_val = _job_title_input
        st.rerun()

    # Display rewrite results
    data = st.session_state.get("rewrite_result")
    if data:
        render_rewrite_results(data)


def _make_resume_docx(resume_text: str, format_style: str = "ats") -> bytes:
    """
    Convert plain-text resume to a formatted DOCX.
    format_style: "ats" | "harvard" | "modern" | "executive"
    """
    import io
    import re
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    fs = format_style or "ats"

    # ── Format theme definitions ──────────────────────────────────────────────
    THEMES = {
        "ats": dict(
            font="Calibri", body_size=10, name_size=16, header_size=11,
            margin=0.75, header_bold=True, header_caps=True,
            name_center=False, header_underline=False,
            header_color=None, accent_color=None,
            bullet_char="\u2022", bullet_indent=0.25,
            header_sb=8, header_sa=3,
        ),
        "harvard": dict(
            font="Times New Roman", body_size=11, name_size=14, header_size=11,
            margin=1.0, header_bold=True, header_caps=True,
            name_center=True, header_underline=True,
            header_color=None, accent_color=None,
            bullet_char="\u2013", bullet_indent=0.25,
            header_sb=10, header_sa=2,
        ),
        "modern": dict(
            font="Calibri", body_size=10, name_size=18, header_size=10,
            margin=0.75, header_bold=True, header_caps=True,
            name_center=False, header_underline=False,
            header_color=(30, 58, 138), accent_color=(30, 58, 138),
            bullet_char="\u25cf", bullet_indent=0.2,
            header_sb=10, header_sa=2,
        ),
        "executive": dict(
            font="Georgia", body_size=10.5, name_size=16, header_size=11,
            margin=1.0, header_bold=True, header_caps=True,
            name_center=True, header_underline=False,
            header_color=None, accent_color=None,
            bullet_char="\u2022", bullet_indent=0.3,
            header_sb=10, header_sa=3,
        ),
    }
    t = THEMES.get(fs, THEMES["ats"])

    doc = Document()
    for sec in doc.sections:
        m = Inches(t["margin"])
        sec.top_margin = sec.bottom_margin = m
        sec.left_margin = sec.right_margin = m

    normal = doc.styles["Normal"]
    normal.font.name = t["font"]
    normal.font.size = Pt(t["body_size"])

    def _strip_md(text):
        return re.sub(r"\*\*(.*?)\*\*", r"\1", text)

    def _add_run(p, text, bold=False, size=None, color=None, underline=False, italic=False):
        run = p.add_run(_strip_md(text))
        run.bold = bold
        run.italic = italic
        run.underline = underline
        run.font.name = t["font"]
        run.font.size = Pt(size or t["body_size"])
        if color:
            run.font.color.rgb = RGBColor(*color)
        return run

    def _para(text, bold=False, size=None, sb=0, sa=2,
              center=False, color=None, underline=False, italic=False):
        p = doc.add_paragraph()
        if center:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _add_run(p, text, bold=bold, size=size, color=color, underline=underline, italic=italic)
        p.paragraph_format.space_before = Pt(sb)
        p.paragraph_format.space_after = Pt(sa)
        return p

    def _section_header(text):
        """Render a section header using the current theme."""
        p = doc.add_paragraph()
        if t["name_center"]:           # harvard / executive center headers too
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = _add_run(
            p, text,
            bold=t["header_bold"],
            size=t["header_size"],
            color=t["header_color"],
            underline=t["header_underline"],
        )
        p.paragraph_format.space_before = Pt(t["header_sb"])
        p.paragraph_format.space_after  = Pt(t["header_sa"])

        # Modern: add a thin colored rule below the header
        if fs == "modern":
            _hrule(color="1e3a8a", sz=4)
        return p

    def _bullet(text):
        text = _strip_md(text.lstrip("\u2022\u25cf\u2013-").strip())
        p = doc.add_paragraph()
        run = p.add_run(f"{t['bullet_char']}  {text}")
        run.font.name = t["font"]
        run.font.size = Pt(t["body_size"])
        pf = p.paragraph_format
        pf.left_indent = Inches(t["bullet_indent"])
        pf.space_before = Pt(0)
        pf.space_after  = Pt(1.5 if fs == "executive" else 1)

    def _hrule(color="94a3b8", sz=6):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after  = Pt(3)
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), str(sz))
        bottom.set(qn("w:space"), "1")
        bottom.set(qn("w:color"), color)
        pBdr.append(bottom)
        pPr.append(pBdr)

    DATE_RE = re.compile(
        r"^(Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*[\s,]+\d{4}"
        r"|^\d{4}\s*[-\u2013]\s*(\d{4}|Present)",
        re.IGNORECASE,
    )

    lines = resume_text.strip().splitlines()
    name_done = contact_done = False

    for line in lines:
        s = line.strip()

        # Separator line (___ or ---)
        if s and len(s) > 5 and len(set(s) - {" "}) == 1 and s[0] in "_-=─":
            if fs != "modern":   # modern inserts its own rule after headers
                _hrule()
            continue

        if not s:
            continue

        # First non-empty line → Name
        if not name_done:
            p = _para(
                s, bold=True, size=t["name_size"], sb=0, sa=2,
                center=t["name_center"],
                color=t["accent_color"] if fs == "modern" else None,
            )
            # Modern: accent line under name
            if fs == "modern":
                _hrule(color="1e3a8a", sz=8)
            name_done = True
            continue

        # Second non-empty line → Contact info
        if not contact_done:
            _para(s, bold=False, size=t["body_size"] - 0.5,
                  sb=0, sa=6, center=t["name_center"])
            contact_done = True
            continue

        # Section header: ALL CAPS, no pipe/bullet
        if (
            s.isupper() and 3 <= len(s) <= 70
            and "\u2022" not in s and "\u25cf" not in s
            and "|" not in s and not s[0].isdigit()
        ):
            _section_header(s)
            continue

        # Bullet point
        if s.startswith(("\u2022", "\u25cf", "\u2013")) or (
            s.startswith("-") and len(s) > 3 and s[1] == " "
        ):
            _bullet(s)
            continue

        # Job/company line with pipe separator
        if "|" in s and not s.isupper():
            _para(s, bold=True, size=t["body_size"], sb=6, sa=1,
                  color=t["accent_color"] if fs == "modern" else None)
            continue

        # Date line
        if DATE_RE.search(s):
            _para(s, bold=False, size=t["body_size"] - 0.5, sb=0, sa=1,
                  italic=(fs == "harvard"))
            continue

        # Regular text
        _para(s, bold=False, size=t["body_size"], sb=1, sa=2)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def render_rewrite_results(data: dict):
    """Show before/after scores, changes, and download button."""
    original = data.get("original_scores", {})
    rewritten = data.get("rewritten_scores", {})

    # Company / job title stored when rewrite was triggered
    company = st.session_state.get("rewrite_company_val", "")
    job_title = st.session_state.get("rewrite_job_title_val", "")

    st.markdown("---")

    # ─── Before / After score comparison ──────────────────────────────────
    st.markdown("##### Score Comparison")

    llm_ats = rewritten.get("llm_ats", 0)
    llm_hr = rewritten.get("llm_hr", 0)
    has_llm = llm_ats > 0 or llm_hr > 0

    score_col1, score_col2, score_col3, score_col4 = st.columns(4)
    with score_col1:
        st.metric("Original ATS", f"{original.get('ats', 0):.0f}%")
    with score_col2:
        ats_delta = rewritten.get("ats", 0) - original.get("ats", 0)
        st.metric("Rewritten ATS", f"{rewritten.get('ats', 0):.0f}%", delta=f"{ats_delta:+.0f}%")
    with score_col3:
        st.metric("Original HR", f"{original.get('hr', 0):.0f}%")
    with score_col4:
        hr_delta = rewritten.get("hr", 0) - original.get("hr", 0)
        st.metric("Rewritten HR", f"{rewritten.get('hr', 0):.0f}%", delta=f"{hr_delta:+.0f}%")

    if has_llm:
        llm_col1, llm_col2, _ = st.columns(3)
        with llm_col1:
            st.metric("LLM ATS", f"{llm_ats:.0f}%", help="Claude's independent ATS assessment")
        with llm_col2:
            st.metric("LLM HR", f"{llm_hr:.0f}%", help="Claude's independent HR recruiter assessment")

    # Visual gauge comparison
    gauge_cols = st.columns(3) if has_llm else st.columns(2)
    with gauge_cols[0]:
        st.plotly_chart(make_gauge(rewritten.get("ats", 0), "Rewritten ATS Score"), use_container_width=True)
    with gauge_cols[1]:
        st.plotly_chart(make_gauge(rewritten.get("hr", 0), "Rewritten HR Score"), use_container_width=True)
    if has_llm:
        with gauge_cols[2]:
            llm_avg = (llm_ats + llm_hr) / 2
            st.plotly_chart(make_gauge(llm_avg, "LLM Overall Score"), use_container_width=True)

    # ─── Explanation ──────────────────────────────────────────────────────
    explanation = data.get("explanation", "")
    if explanation:
        st.markdown("##### Tailoring Strategy")
        st.info(explanation)

    # ─── Changes made ─────────────────────────────────────────────────────
    changes = data.get("changes_made", [])
    if changes:
        st.markdown("##### Changes Made")
        for change in changes:
            st.markdown(f"- {change}")

    # ─── Rewritten resume text + DOCX download ────────────────────────────
    rewritten_text = data.get("rewritten_resume", "")
    if rewritten_text:
        st.markdown("---")
        st.markdown("##### Tailored Resume")
        st.text_area("Preview your tailored resume", value=rewritten_text, height=400, key="rewritten_output")

        fmt_style = data.get("format_style", "ats")
        fmt_labels = {
            "ats": "ATS Standard", "harvard": "Harvard Classic",
            "modern": "Modern Clean", "executive": "Executive Classic",
        }
        fmt_label = fmt_labels.get(fmt_style, "ATS Standard")

        # Build a descriptive filename: Company_JobTitle_Resume_style.docx
        def _safe(s: str) -> str:
            import re
            return re.sub(r"[^\w\-]", "_", s.strip()).strip("_") if s.strip() else ""

        _name_parts = [p for p in [_safe(company), _safe(job_title), "Resume", fmt_style] if p]
        _docx_name = "_".join(_name_parts) + ".docx"
        _txt_name = "_".join(_name_parts[:-1] + ["Resume"]) + ".txt" if _name_parts else "tailored_resume.txt"

        dl_col, _ = st.columns([2, 3])
        with dl_col:
            try:
                docx_bytes = _make_resume_docx(rewritten_text, format_style=fmt_style)
                st.download_button(
                    label=f"Download — {fmt_label} (.docx)",
                    data=docx_bytes,
                    file_name=_docx_name,
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True,
                    type="primary",
                )
            except Exception:
                st.download_button(
                    label="Download Tailored Resume (.txt)",
                    data=rewritten_text,
                    file_name=_txt_name,
                    mime="text/plain",
                    use_container_width=True,
                    type="primary",
                )

    # ─── Auto-add to tracker (once per rewrite result) ────────────────────
    if is_authenticated() and company and rewritten_text:
        _tracker_key = f"tracker_saved_{id(data)}"
        if not st.session_state.get(_tracker_key):
            _resume_fname = _docx_name if rewritten_text else ""
            _tracker_resp = api(
                "POST", "/tracker/add",
                {
                    "company": company,
                    "job_title": job_title,
                    "status": "Applied",
                    "resume_file": _resume_fname,
                    "ats_score": rewritten.get("ats", 0),
                    "hr_score": rewritten.get("hr", 0),
                    "llm_score": rewritten.get("llm_hr", 0),
                },
                token=st.session_state.token,
            )
            if _tracker_resp.get("status") in (200, 201) or _tracker_resp.get("data", {}).get("id"):
                st.session_state[_tracker_key] = True
                st.success(f"Saved to Job Tracker — {company} · {job_title}")

    st.markdown(f"**Model:** {data.get('model_used', 'claude-sonnet-4-6')}")


# ─── Auth pages ──────────────────────────────────────────────────────────────

def page_register():
    """Registration page."""
    _, form_col, _ = st.columns([1, 2, 1])
    with form_col:
        st.markdown('<div class="section-title">Create Account</div>', unsafe_allow_html=True)
        st.markdown(
            '<div class="section-subtitle">Sign up to track your usage and unlock Pro features.</div>',
            unsafe_allow_html=True,
        )

        with st.form("register_form"):
            email = st.text_input("Email", placeholder="you@example.com")
            password = st.text_input("Password", type="password", placeholder="Min 6 characters")
            submitted = st.form_submit_button("Create Account", use_container_width=True, type="primary")

        if submitted:
            if not email or not password:
                st.error("Email and password are required.")
                return
            if len(password) < 6:
                st.error("Password must be at least 6 characters.")
                return

            with st.spinner("Creating account..."):
                result = api("POST", "/auth/register", {"email": email, "password": password})

            if result["status"] == 200:
                data = result["data"]
                st.session_state.token = data["token"]
                st.session_state.user = data["user"]
                st.session_state.page = "scorer"
                # Persist auth in URL param (survives page refresh)
                st.query_params["_t"] = data["token"]
                # If user had a resume before registering, save it to their account now
                pre_existing = st.session_state.get("stored_resume", "")
                if pre_existing:
                    if _upload_resume_to_cloud(data["token"], pre_existing):
                        st.session_state.resume_on_file = True
                st.success("Account created! Redirecting to scorer...")
                st.rerun()
            elif result["status"] == 409:
                st.error("Email already registered. Try logging in.")
            else:
                st.error(result["data"].get("detail", "Registration failed."))

        st.markdown("---")
        st.markdown("Already have an account?")
        if st.button("Log in instead"):
            st.session_state.page = "login"
            st.rerun()


def page_login():
    """Login page."""
    _, form_col, _ = st.columns([1, 2, 1])
    with form_col:
        st.markdown('<div class="section-title">Welcome Back</div>', unsafe_allow_html=True)
        st.markdown(
            '<div class="section-subtitle">Sign in to access your dashboard and scoring history.</div>',
            unsafe_allow_html=True,
        )

        with st.form("login_form"):
            email = st.text_input("Email", placeholder="you@example.com")
            password = st.text_input("Password", type="password")
            submitted = st.form_submit_button("Sign In", use_container_width=True, type="primary")

        if submitted:
            if not email or not password:
                st.error("Email and password are required.")
                return

            with st.spinner("Signing in..."):
                result = api("POST", "/auth/login", {"email": email, "password": password})

            if result["status"] == 200:
                data = result["data"]
                st.session_state.token = data["token"]
                st.session_state.user = data["user"]
                st.session_state.page = "dashboard"
                # Persist auth in URL param (survives page refresh)
                st.query_params["_t"] = data["token"]
                # Auto-load saved resume from cloud into session
                saved = _fetch_saved_resume(data["token"])
                if saved:
                    st.session_state.stored_resume = saved
                    st.session_state.resume_on_file = True
                st.success("Signed in!")
                st.rerun()
            else:
                st.error(result["data"].get("detail", "Login failed. Check your credentials."))

        st.markdown("---")
        st.markdown("Don't have an account?")
        if st.button("Create one now"):
            st.session_state.page = "register"
            st.rerun()


# ─── Dashboard ───────────────────────────────────────────────────────────────

def page_dashboard():
    """Authenticated user dashboard."""
    if not is_authenticated():
        st.session_state.page = "login"
        st.rerun()
        return

    user = st.session_state.user
    token = st.session_state.token

    st.markdown('<div class="section-title">Dashboard</div>', unsafe_allow_html=True)
    st.markdown(
        f'<div class="section-subtitle">Signed in as {user["email"]} &middot; '
        f'{user.get("tier", "free").title()} plan</div>',
        unsafe_allow_html=True,
    )

    # Fetch usage stats
    usage = api("GET", "/auth/usage", token=token)

    if usage["status"] == 200:
        stats = usage["data"]
        total_used = stats.get("total_scores", 0)
        today_used = stats.get("today_scores", 0)
        remaining = stats.get("remaining")
        tier = user.get("tier", "free")

        # Metrics row
        rewrites_info = stats.get("rewrites", {})
        if tier in ("pro", "ultra"):
            m1, m2, m3, m4 = st.columns(4)
        else:
            m1, m2, m3 = st.columns(3)
        with m1:
            st.metric("Total Scores", total_used)
        with m2:
            st.metric("Today", today_used)
        with m3:
            if tier == "ultra":
                st.metric("Plan", "Ultra")
            elif tier == "pro":
                st.metric("Plan", "Pro")
            elif remaining is not None:
                st.metric("Remaining", f"{remaining}/5")
            else:
                st.metric("Remaining", "5")
        if tier in ("pro", "ultra"):
            with m4:
                if tier == "ultra":
                    st.metric("Rewrites", "Unlimited")
                else:
                    rw_remaining = rewrites_info.get("remaining", 10)
                    rw_limit = rewrites_info.get("limit", 10)
                    st.metric("Rewrites", f"{rw_remaining}/{rw_limit}")

        # Usage progress bar (free tier)
        if tier == "free" and remaining is not None:
            used_pct = min(1.0, total_used / 5)
            st.progress(used_pct, text=f"{total_used}/5 free scores used")
    elif usage["status"] == 401:
        st.warning("Session expired. Please log in again.")
        st.query_params.clear()
        st.session_state.token = None
        st.session_state.user = None
        st.session_state.page = "login"
        st.rerun()
        return
    else:
        st.warning("Could not fetch usage stats.")

    st.markdown("---")

    # Upgrade / Billing section
    tier = user.get("tier", "free")
    if tier == "free":
        up1, up2 = st.columns(2)
        with up1:
            st.markdown("""
            <div class="card-accent">
                <p style="color: #818cf8; font-weight: 700; font-size: 16px; margin-bottom: 4px;">Pro — $12/month</p>
                <p style="color: #94a3b8; font-size: 13px;">Unlimited scoring + LLM analysis + 10 AI rewrites/month</p>
            </div>
            """, unsafe_allow_html=True)
            if st.button("Upgrade to Pro", use_container_width=True, type="primary", key="dash_pro"):
                with st.spinner("Creating checkout session..."):
                    result = api("POST", "/billing/checkout", {"tier": "pro"}, token=token)
                if result["status"] == 200 and "checkout_url" in result["data"]:
                    st.link_button("Pay on Stripe", result["data"]["checkout_url"], use_container_width=True)
                elif result["status"] == 503:
                    st.warning("Stripe billing is not configured yet.")
                else:
                    st.error(result["data"].get("detail", "Checkout failed."))
        with up2:
            st.markdown("""
            <div class="card-accent" style="border-color: #a855f7;">
                <p style="color: #a855f7; font-weight: 700; font-size: 16px; margin-bottom: 4px;">Ultra — $29/month</p>
                <p style="color: #94a3b8; font-size: 13px;">Everything in Pro + unlimited AI rewrites</p>
            </div>
            """, unsafe_allow_html=True)
            if st.button("Upgrade to Ultra", use_container_width=True, type="primary", key="dash_ultra"):
                with st.spinner("Creating checkout session..."):
                    result = api("POST", "/billing/checkout", {"tier": "ultra"}, token=token)
                if result["status"] == 200 and "checkout_url" in result["data"]:
                    st.link_button("Pay on Stripe", result["data"]["checkout_url"], use_container_width=True)
                elif result["status"] == 503:
                    st.warning("Stripe billing is not configured yet.")
                else:
                    st.error(result["data"].get("detail", "Checkout failed."))

    elif tier == "pro":
        st.markdown("""
        <div class="card-accent" style="border-color: #a855f7;">
            <p style="color: #a855f7; font-weight: 700; font-size: 16px; margin-bottom: 4px;">Upgrade to Ultra — $29/month</p>
            <p style="color: #94a3b8; font-size: 13px;">Everything you have now + unlimited AI rewrites (instead of 10/month)</p>
        </div>
        """, unsafe_allow_html=True)
        if st.button("Upgrade to Ultra", use_container_width=True, type="primary", key="dash_pro_to_ultra"):
            with st.spinner("Creating checkout session..."):
                result = api("POST", "/billing/checkout", {"tier": "ultra"}, token=token)
            if result["status"] == 200 and "checkout_url" in result["data"]:
                st.link_button("Pay on Stripe", result["data"]["checkout_url"], use_container_width=True)
            elif result["status"] == 503:
                st.warning("Stripe billing is not configured yet.")
            else:
                st.error(result["data"].get("detail", "Checkout failed."))
    else:
        st.markdown("##### Manage Subscription")
        if st.button("Open Billing Portal", use_container_width=True):
            with st.spinner("Opening billing portal..."):
                result = api("POST", "/billing/portal", token=token)

            if result["status"] == 200:
                url = result["data"] if isinstance(result["data"], str) else result["data"].get("url", "")
                if url:
                    st.link_button("Open Stripe Billing Portal", url, use_container_width=True)
                else:
                    st.warning("Could not retrieve portal URL.")
            else:
                st.error(result["data"].get("detail", "Billing portal unavailable."))

    # ─── Claude Code Plugin setup (Pro / Ultra only) ─────────────────────
    if tier in ("pro", "ultra"):
        st.markdown("---")
        st.markdown("##### Claude Code Plugin Setup")
        st.markdown(
            '<p style="color: #94a3b8; font-size: 13px;">'
            "Generate an API key to connect your account to the local Claude Code plugin. "
            "<strong>Claude Code users only need Pro</strong> — your Anthropic subscription "
            "handles resume writing; this key unlocks unlimited <em>scoring</em> in the plugin."
            "</p>",
            unsafe_allow_html=True,
        )

        with st.form("plugin_apikey_form"):
            label = st.text_input("Key label (optional)", placeholder="e.g. my-laptop")
            create_key = st.form_submit_button("Generate Plugin API Key", type="primary")

        if create_key:
            with st.spinner("Generating..."):
                result = api("POST", "/auth/api-key", {"label": label or "plugin"}, token=token)

            if result["status"] == 200:
                api_key = result["data"]["api_key"]
                st.success("API key created. Copy it now — it won't be shown again.")
                st.code(api_key, language=None)
                st.markdown("**Add to your `.env` file in the Resume Builder folder:**")
                st.code(f"SCORER_CLOUD_URL=https://resume-scorer.fly.dev\nSCORER_CLOUD_API_KEY={api_key}", language="bash")
                st.markdown(
                    '<p style="color: #94a3b8; font-size: 13px; margin-top: 8px;">'
                    "After saving, restart the scorer server: "
                    "<code>powershell -ExecutionPolicy Bypass -File restart_scorer.ps1</code>"
                    "</p>",
                    unsafe_allow_html=True,
                )
            else:
                st.error(result["data"].get("detail", "Failed to create API key."))



# ─── Cover Letter page ───────────────────────────────────────────────────────


def page_cover_letter():
    """Pro + Ultra tier: AI cover letter generation."""
    st.markdown('<div class="section-title">AI Cover Letter</div>', unsafe_allow_html=True)
    st.markdown(
        '<div class="section-subtitle">Generate a tailored cover letter from your resume and a job description.</div>',
        unsafe_allow_html=True,
    )

    user_tier = ""
    if is_authenticated() and st.session_state.user:
        user_tier = st.session_state.user.get("tier", "free")

    can_generate = user_tier in ("pro", "ultra")

    if not can_generate:
        st.markdown("""
        <div class="card-accent" style="text-align: center; border-color: #818cf8;">
            <p style="color: #818cf8; font-weight: 700; font-size: 20px; margin-bottom: 8px;">
                Pro Feature
            </p>
            <p style="color: #94a3b8; font-size: 15px; margin-bottom: 16px;">
                AI cover letter generation is available on Pro ($12/month) and
                Ultra ($29/month). Claude AI writes a compelling, ready-to-send
                cover letter tailored to each job.
            </p>
            <p style="color: #94a3b8; font-size: 14px;">
                &#10003; Personalized to each JD &nbsp;&nbsp;
                &#10003; Highlights your best matches &nbsp;&nbsp;
                &#10003; Ready to send (no blanks)
            </p>
        </div>
        """, unsafe_allow_html=True)

        st.markdown("")
        if not is_authenticated():
            _, cta_col, _ = st.columns([2, 3, 2])
            with cta_col:
                if st.button("Sign Up to Get Started", type="primary", use_container_width=True, key="cl_signup"):
                    st.session_state.page = "register"
                    st.rerun()
        else:
            _, pro_col, ultra_col, _ = st.columns([1, 2, 2, 1])
            with pro_col:
                if st.button("Pro — $12/month", type="primary", use_container_width=True, key="cl_pro"):
                    with st.spinner("Creating checkout session..."):
                        result = api("POST", "/billing/checkout", {"tier": "pro"}, token=st.session_state.token)
                    if result["status"] == 200 and "checkout_url" in result["data"]:
                        st.link_button("Complete Payment on Stripe", result["data"]["checkout_url"], use_container_width=True)
                    elif result["status"] == 503:
                        st.warning("Stripe billing is not configured yet.")
                    else:
                        st.error(result["data"].get("detail", "Could not create checkout session."))
            with ultra_col:
                if st.button("Ultra — $29/month", use_container_width=True, key="cl_ultra"):
                    with st.spinner("Creating checkout session..."):
                        result = api("POST", "/billing/checkout", {"tier": "ultra"}, token=st.session_state.token)
                    if result["status"] == 200 and "checkout_url" in result["data"]:
                        st.link_button("Complete Payment on Stripe", result["data"]["checkout_url"], use_container_width=True)
                    elif result["status"] == 503:
                        st.warning("Stripe billing is not configured yet.")
                    else:
                        st.error(result["data"].get("detail", "Could not create checkout session."))
        return

    # ─── Pro / Ultra user: show the generator ─────────────────────────────
    prefill_resume = st.session_state.pop("prefill_resume", "") or ""
    prefill_jd = st.session_state.pop("prefill_jd", "") or ""
    prefill_job_url = st.session_state.pop("prefill_job_url", "") or ""
    prefill_job_title = st.session_state.pop("prefill_job_title", "") or ""
    if prefill_jd or prefill_job_url:
        prefill_jd, jd_warn = _apply_jd_prefill("cover_jd_text", prefill_jd, prefill_job_url, prefill_job_title)
        if jd_warn:
            st.warning(f"⚠ {jd_warn}")

    col_resume, col_jd = st.columns(2)
    with col_resume:
        resume_text = resume_input("Your Resume", prefill=prefill_resume, key_prefix="cover", height=350)
    with col_jd:
        jd_text = st.text_area(
            "Target Job Description",
            height=350,
            placeholder="Paste the job description here...",
            key="cover_jd_text",
        )

    _, btn_col, _ = st.columns([3, 2, 3])
    with btn_col:
        generate_clicked = st.button("Generate Cover Letter", use_container_width=True, type="primary")

    if generate_clicked:
        if not resume_text or not jd_text:
            st.error("Please paste both your resume and the job description.")
            return
        if len(resume_text.strip()) < 100:
            st.warning("Resume seems too short. Please paste the full text.")
            return

        with st.spinner("Claude is writing your cover letter... This may take 15-30 seconds."):
            result = api("POST", "/cover-letter", {
                "resume_text": resume_text,
                "jd_text": jd_text,
            }, token=st.session_state.token)

        if result["status"] != 200:
            st.error(f"Generation failed: {result['data'].get('detail', 'Unknown error')}")
            return

        st.session_state.cover_letter_result = result["data"]
        st.session_state.cover_letter_resume_text = resume_text

    # Display result
    data = st.session_state.get("cover_letter_result")
    if data:
        render_cover_letter_result(data)


def _extract_jd_meta(jd_text: str):
    """Extract job title and company from JD text using lightweight heuristics."""
    import re
    lines = [l.strip() for l in jd_text.splitlines() if l.strip()]
    title, company = "", ""
    # Job title: first non-URL line under 80 chars that doesn't end in a period
    for line in lines[:8]:
        if len(line) < 80 and not line.startswith("http") and not line.endswith("."):
            title = line
            break
    # Company: look for "Company: X", "About X", "Employer: X" patterns
    for line in lines[:30]:
        m = re.match(r'(?:company|employer|organization)[:\s-]+(.+)', line, re.I)
        if m:
            company = m.group(1).strip()[:60]
            break
        m = re.match(r'about\s+([\w\s&,\.\-]+?)(?:\s*[-|]|$)', line, re.I)
        if m:
            candidate = m.group(1).strip()
            if 2 < len(candidate) < 60 and not any(
                w in candidate.lower() for w in ["us ", "the role", "this role", "you "]
            ):
                company = candidate
                break
    return title.strip(), company.strip()


def _safe_filename(s: str, maxlen: int = 30) -> str:
    """Sanitise a string for use in a filename."""
    import re
    return re.sub(r'[^\w\-]', '_', s).strip('_')[:maxlen]


def _parse_sender_info(resume_text: str):
    """Extract sender name and contact line from the top of a plain-text resume."""
    import re
    lines = [l.strip() for l in resume_text.splitlines() if l.strip()]
    name = ""
    contact = ""
    for i, line in enumerate(lines[:5]):
        # First non-empty line is the name — strip markdown # prefix if present
        if not name:
            name = re.sub(r"^#+\s*", "", line).strip()
            continue
        # Second line: looks like contact info (contains | or @ or phone digits)
        if re.search(r"[@|]|\d{3}[-.\s]\d{3}", line):
            contact = line
            break
    return name, contact


def _make_cover_letter_docx(
    paragraphs: list,
    sender_name: str,
    sender_contact: str,
    company: str,
    job_title: str,
) -> bytes:
    """Build a one-page professional cover letter DOCX."""
    import io
    from datetime import date
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    # ── Page setup: 1" margins, Letter size ──────────────────────────────────
    for section in doc.sections:
        section.top_margin    = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin   = Inches(1.1)
        section.right_margin  = Inches(1.1)

    FONT = "Calibri"
    BODY_PT = 11

    def _run(para, text, bold=False, size=BODY_PT, color=None):
        run = para.add_run(text)
        run.bold = bold
        run.font.name = FONT
        run.font.size = Pt(size)
        if color:
            run.font.color.rgb = RGBColor(*color)
        return run

    def _blank(space_after=4):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(space_after)

    def _body(text, space_after=10):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(space_after)
        _run(p, text)
        return p

    # ── Sender block (top-left) ───────────────────────────────────────────────
    name_para = doc.add_paragraph()
    name_para.paragraph_format.space_before = Pt(0)
    name_para.paragraph_format.space_after  = Pt(2)
    _run(name_para, sender_name or "Applicant Name", bold=True, size=13)

    if sender_contact:
        contact_para = doc.add_paragraph()
        contact_para.paragraph_format.space_before = Pt(0)
        contact_para.paragraph_format.space_after  = Pt(0)
        _run(contact_para, sender_contact, size=10, color=(71, 85, 105))

    _blank(space_after=8)

    # ── Date ─────────────────────────────────────────────────────────────────
    today = date.today().strftime("%B %d, %Y")
    _body(today, space_after=8)

    # ── Recipient block ───────────────────────────────────────────────────────
    _body("Hiring Manager")
    if company:
        _body(company, space_after=8)
    else:
        _blank(space_after=8)

    # ── Salutation ────────────────────────────────────────────────────────────
    _body("Dear Hiring Manager,", space_after=10)

    # ── Body paragraphs ───────────────────────────────────────────────────────
    for para_text in paragraphs:
        if para_text.strip():
            _body(para_text.strip(), space_after=10)

    # ── Closing ───────────────────────────────────────────────────────────────
    _blank(space_after=4)
    _body("Sincerely,", space_after=40)   # space for handwritten signature
    _body(sender_name or "Applicant Name", space_after=2)
    if job_title:
        _run(doc.add_paragraph(), job_title, size=10, color=(71, 85, 105))

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


def render_cover_letter_result(data: dict):
    """Display generated cover letter with copy/download options."""
    paragraphs = data.get("paragraphs", [])
    full_text = data.get("full_text", "")
    company = data.get("company", "")
    title = data.get("job_title", "")
    word_count = data.get("word_count", 0)

    st.markdown("---")

    # Header info
    meta_cols = st.columns(3)
    with meta_cols[0]:
        st.metric("Company", company or "Detected")
    with meta_cols[1]:
        st.metric("Position", title or "Detected")
    with meta_cols[2]:
        color = "#22c55e" if word_count <= 400 else "#eab308"
        st.metric("Word Count", f"{word_count}")
        if word_count > 400:
            st.caption("Slightly over 1-page target")

    st.markdown("---")

    # Render each paragraph
    st.markdown("### Your Cover Letter")
    for i, para in enumerate(paragraphs):
        st.markdown(f"<p style='color: #e2e8f0; line-height: 1.7; margin-bottom: 16px;'>{para}</p>", unsafe_allow_html=True)

    # Copy/download
    st.markdown("---")
    resume_text_for_cl = st.session_state.get("cover_letter_resume_text", "")
    sender_name, sender_contact = _parse_sender_info(resume_text_for_cl)

    try:
        docx_bytes = _make_cover_letter_docx(
            paragraphs=paragraphs,
            sender_name=sender_name,
            sender_contact=sender_contact,
            company=company,
            job_title=title,
        )
        import re as _re
        def _cl_safe(s: str) -> str:
            return _re.sub(r"[^\w\-]", "_", s.strip()).strip("_") if s and s.strip() else ""
        _cl_parts = [p for p in [_cl_safe(company), _cl_safe(title), "Cover_Letter"] if p]
        _cl_base = "_".join(_cl_parts) if _cl_parts else "Cover_Letter"
        action_cols = st.columns([1, 1, 2])
        with action_cols[0]:
            st.download_button(
                "Download as .docx",
                data=docx_bytes,
                file_name=f"{_cl_base}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True,
                type="primary",
            )
        with action_cols[1]:
            st.download_button(
                "Download as .txt",
                data=full_text,
                file_name=f"{_cl_base}.txt",
                mime="text/plain",
                use_container_width=True,
            )
    except Exception:
        import re as _re
        _fb = _re.sub(r"[^\w\-]", "_", company.strip()).strip("_") if company and company.strip() else "CoverLetter"
        action_cols = st.columns([1, 3])
        with action_cols[0]:
            st.download_button(
                "Download as .txt",
                data=full_text,
                file_name=f"{_fb}_Cover_Letter.txt",
                mime="text/plain",
                use_container_width=True,
            )


# ─── Job Discovery page ──────────────────────────────────────────────────────


def render_job_card(rank: int, job: dict, resume_text: str = ""):
    """Render a single job result card with action buttons."""
    ats = job.get("ats_score", 0)
    hr = job.get("hr_score", 0)

    # Score color
    if ats >= 75:
        ats_color = "#22c55e"
    elif ats >= 60:
        ats_color = "#eab308"
    else:
        ats_color = "#ef4444"

    if hr >= 70:
        hr_color = "#22c55e"
    elif hr >= 55:
        hr_color = "#eab308"
    else:
        hr_color = "#ef4444"

    salary_text = ""
    if job.get("salary_min") and job.get("salary_max"):
        salary_text = f"${job['salary_min']:,.0f} – ${job['salary_max']:,.0f}"
    elif job.get("salary_min"):
        salary_text = f"${job['salary_min']:,.0f}+"

    rec = job.get("hr_detail", {}).get("recommendation", "")
    rec_badge = ""
    if rec:
        badge_colors = {
            "STRONG INTERVIEW": "#22c55e",
            "INTERVIEW": "#22c55e",
            "MAYBE": "#eab308",
            "PASS": "#ef4444",
        }
        bc = badge_colors.get(rec, "#94a3b8")
        rec_badge = f'<span style="background: {bc}22; color: {bc}; padding: 2px 8px; border-radius: 4px; font-size: 12px; font-weight: 600;">{rec}</span>'

    import html as _html
    matched_kw = job.get("ats_detail", {}).get("matched_keywords", [])
    missing_kw = job.get("ats_detail", {}).get("missing_keywords", [])

    matched_chips = " ".join(
        f'<span style="background: #22c55e22; color: #22c55e; padding: 2px 6px; border-radius: 4px; font-size: 11px; margin: 1px;">{_html.escape(str(kw))}</span>'
        for kw in matched_kw[:8]
    )
    missing_chips = " ".join(
        f'<span style="background: #ef444422; color: #ef4444; padding: 2px 6px; border-radius: 4px; font-size: 11px; margin: 1px;">{_html.escape(str(kw))}</span>'
        for kw in missing_kw[:5]
    )

    title_esc = _html.escape(job.get("title", "Unknown"))
    company_esc = _html.escape(job.get("company", ""))
    location_esc = _html.escape(job.get("location", ""))
    posted = job.get("posted_date", "")

    salary_html = f' &nbsp;&bull;&nbsp; <span style="color: #818cf8;">{salary_text}</span>' if salary_text else ""
    posted_html = f" &nbsp;&bull;&nbsp; Posted: {posted}" if posted else ""

    card_html = (
        '<div class="card" style="margin-bottom: 12px;">'
        '<div style="display: flex; justify-content: space-between; align-items: start;">'
        '<div>'
        f'<span style="color: #818cf8; font-weight: 700; font-size: 14px;">#{rank}</span>'
        f'<span style="color: #e2e8f0; font-weight: 600; font-size: 16px; margin-left: 8px;">{title_esc}</span>'
        f'{rec_badge}'
        '</div>'
        '<div style="text-align: right;">'
        f'<span style="color: {ats_color}; font-weight: 700; font-size: 18px;">ATS {ats}%</span>'
        '<span style="color: #475569; margin: 0 6px;">|</span>'
        f'<span style="color: {hr_color}; font-weight: 700; font-size: 18px;">HR {hr}%</span>'
        '</div>'
        '</div>'
        '<div style="margin-top: 6px; color: #94a3b8; font-size: 14px;">'
        f'<strong>{company_esc}</strong> &nbsp;&bull;&nbsp; {location_esc}'
        f'{salary_html}{posted_html}'
        '</div>'
    )
    if matched_chips or missing_chips:
        card_html += '<div style="margin-top: 8px;">'
        if matched_chips:
            card_html += f'<div style="margin-bottom: 4px;"><span style="color: #64748b; font-size: 11px;">MATCHED:</span> {matched_chips}</div>'
        if missing_chips:
            card_html += f'<div><span style="color: #64748b; font-size: 11px;">MISSING:</span> {missing_chips}</div>'
        card_html += '</div>'
    card_html += '</div>'
    st.markdown(card_html, unsafe_allow_html=True)

    # Action buttons row
    btn_cols = st.columns([1, 1, 1, 1, 2])
    jd_text = job.get("description", "")

    with btn_cols[0]:
        if job.get("url"):
            st.link_button("View Listing", job["url"], use_container_width=True)

    # listing_url = Adzuna /details/{id} page (scrapeable); url = employer redirect
    job_url = job.get("listing_url", "") or job.get("url", "")
    job_title_nav = job.get("title", "")
    company_nav = job.get("company", "")

    with btn_cols[1]:
        if st.button("Score", key=f"score_{rank}", use_container_width=True):
            st.session_state.prefill_resume = resume_text
            st.session_state.prefill_jd = jd_text
            st.session_state.prefill_job_url = job_url
            st.session_state.prefill_job_title = job_title_nav
            st.session_state.prefill_company = company_nav
            st.session_state.page = "scorer"
            st.rerun()

    with btn_cols[2]:
        if st.button("Tailor Resume", key=f"tailor_{rank}", use_container_width=True):
            st.session_state.prefill_resume = resume_text
            st.session_state.prefill_jd = jd_text
            st.session_state.prefill_job_url = job_url
            st.session_state.prefill_job_title = job_title_nav
            st.session_state.prefill_company = company_nav
            st.session_state.page = "rewriter"
            st.rerun()

    with btn_cols[3]:
        if st.button("Cover Letter", key=f"cl_{rank}", use_container_width=True, type="primary"):
            st.session_state.prefill_resume = resume_text
            st.session_state.prefill_jd = jd_text
            st.session_state.prefill_job_url = job_url
            st.session_state.prefill_job_title = job_title_nav
            st.session_state.prefill_company = company_nav
            st.session_state.page = "cover_letter"
            st.rerun()


def page_discover():
    """Job Discovery page — search and score jobs against your resume."""
    st.markdown("## Discover Jobs")
    st.markdown(
        '<p style="color: #94a3b8;">Search for jobs and see how your resume scores against each one.</p>',
        unsafe_allow_html=True,
    )

    # ─── Auth gate: require login ─────────────────────────────────────────
    if not is_authenticated():
        st.markdown("""
        <div class="card-accent" style="text-align: center; border-color: #818cf8;">
            <p style="color: #818cf8; font-weight: 700; font-size: 20px; margin-bottom: 8px;">
                Sign Up to Discover Jobs
            </p>
            <p style="color: #94a3b8; font-size: 15px; margin-bottom: 16px;">
                Job discovery searches thousands of openings and scores each one against
                your resume with ATS + HR analysis. Create a free account to get started.
            </p>
            <p style="color: #94a3b8; font-size: 14px;">
                &#10003; Free: 2 discoveries included &nbsp;&nbsp;
                &#10003; Pro: unlimited &nbsp;&nbsp;
                &#10003; Real job listings from Adzuna
            </p>
        </div>
        """, unsafe_allow_html=True)
        st.markdown("")
        _, cta_col, _ = st.columns([2, 3, 2])
        with cta_col:
            if st.button("Sign Up Free", type="primary", use_container_width=True, key="discover_signup"):
                st.session_state.page = "register"
                st.rerun()
        return

    # ─── Usage check for free tier ────────────────────────────────────────
    user_tier = st.session_state.user.get("tier", "free") if st.session_state.user else "free"
    if user_tier == "free":
        remaining = max(0, 5 - st.session_state.scores_used)
        if remaining <= 0:
            st.markdown("""
            <div class="card-accent" style="text-align: center; border-color: #eab308;">
                <p style="color: #eab308; font-weight: 700; font-size: 18px; margin-bottom: 8px;">
                    Free Limit Reached
                </p>
                <p style="color: #94a3b8; font-size: 15px;">
                    You've used all 5 free scores. Each job discovery counts as 1 score.
                    Upgrade to Pro for unlimited discoveries.
                </p>
            </div>
            """, unsafe_allow_html=True)
            st.markdown("")
            _, pro_col, ultra_col, _ = st.columns([1, 2, 2, 1])
            with pro_col:
                if st.button("Pro — $12/month", type="primary", use_container_width=True, key="discover_pro"):
                    with st.spinner("Creating checkout session..."):
                        result = api("POST", "/billing/checkout", {"tier": "pro"}, token=st.session_state.token)
                    if result["status"] == 200 and "checkout_url" in result["data"]:
                        st.link_button("Complete Payment on Stripe", result["data"]["checkout_url"], use_container_width=True)
            with ultra_col:
                if st.button("Ultra — $29/month", use_container_width=True, key="discover_ultra"):
                    with st.spinner("Creating checkout session..."):
                        result = api("POST", "/billing/checkout", {"tier": "ultra"}, token=st.session_state.token)
                    if result["status"] == 200 and "checkout_url" in result["data"]:
                        st.link_button("Complete Payment on Stripe", result["data"]["checkout_url"], use_container_width=True)
            return
        st.markdown(
            f'<div style="text-align: right; color: #94a3b8; font-size: 13px; margin-bottom: 12px;">'
            f'Free scores remaining: <strong style="color: #818cf8;">{remaining}/5</strong> (each discovery = 1 score)</div>',
            unsafe_allow_html=True,
        )

    col1, col2 = st.columns([2, 1])

    with col1:
        resume_text = resume_input("Your Resume", key_prefix="discover", height=200)

    with col2:
        job_title = st.text_input("Job title", placeholder="e.g., Data Scientist", key="discover_title")
        location = st.text_input("Location (optional)", placeholder="e.g., New York", key="discover_location")
        remote_only = st.checkbox("Include remote jobs", key="discover_remote")
        max_results = st.slider("Max results", 3, 20, 10, key="discover_max")

    if st.button("Search & Score", type="primary", use_container_width=True, disabled=not resume_text or not job_title):
        with st.spinner("Searching jobs and scoring matches..."):
            result = api("POST", "/jobs/discover", json_data={
                "resume_text": resume_text,
                "job_title": job_title,
                "location": location,
                "remote_only": remote_only,
                "max_results": max_results,
            }, token=st.session_state.token)

            if result["status"] == 200:
                data = result["data"]
                if data.get("setup_required"):
                    st.info(data.get("message", "API keys required for job discovery."))
                    st.session_state.discover_results = None
                elif not data.get("jobs"):
                    st.warning(data.get("message", "No jobs found. Try a different job title or location."))
                    st.session_state.discover_results = None
                else:
                    st.session_state.discover_results = {
                        "jobs": data["jobs"],
                        "resume_text": resume_text,
                        "attribution": data.get("attribution", ""),
                    }
            else:
                detail = result.get("data", {}).get("detail", "Job discovery failed.")
                st.error(detail)
                st.session_state.discover_results = None

    # Render persisted results (survives reruns for button clicks)
    cached = st.session_state.get("discover_results")
    if cached:
        jobs = cached["jobs"]
        cached_resume = cached.get("resume_text", "")
        st.success(f"Found {len(jobs)} matching jobs, ranked by fit score.")

        for job in jobs:
            render_job_card(job["rank"], job, resume_text=cached_resume)

        attr = cached.get("attribution", "")
        if attr:
            st.markdown(
                f'<div style="text-align: center; color: #64748b; font-size: 12px; margin-top: 16px;">{attr}</div>',
                unsafe_allow_html=True,
            )


# ─── Job Tracker page ────────────────────────────────────────────────────────

_TRACKER_STATUSES = ["Applied", "Waiting", "Interview Scheduled", "Interviewed", "Offer", "Rejected", "Withdrawn"]


def page_tracker():
    """Job application tracker — requires login."""
    st.markdown('<div class="section-title">Job Application Tracker</div>', unsafe_allow_html=True)

    if not is_authenticated():
        st.warning("Please sign in to view your job tracker.")
        if st.button("Sign In", type="primary"):
            st.session_state.page = "login"
            st.rerun()
        return

    result = api("GET", "/tracker", token=st.session_state.token)
    if result["status"] != 200:
        _detail = result.get("data", {}).get("detail", f"HTTP {result['status']}")
        st.error(f"Tracker error: {_detail}")
        return

    apps = result.get("data", {}).get("applications", [])

    if not apps:
        st.info("No applications tracked yet. Rewrite a resume and it will appear here automatically.")
        return

    st.markdown(f"**{len(apps)} application{'s' if len(apps) != 1 else ''}** tracked")
    st.markdown("---")

    # Header row
    _hcols = st.columns([2, 2, 1.4, 1, 1, 1, 1.8])
    for _hc, _hl in zip(_hcols, ["Company", "Job Title", "Date", "ATS", "HR", "LLM", "Status"]):
        _hc.markdown(f"**{_hl}**")
    st.markdown("---")

    for app in apps:
        _entry_id = app["id"]
        _company = app.get("company", "")
        _title = app.get("job_title", "")
        _date = (app.get("created_at") or "")[:10]
        _ats = app.get("ats_score") or 0
        _hr = app.get("hr_score") or 0
        _llm = app.get("llm_score") or 0
        _status = app.get("status", "Applied")

        _cols = st.columns([2, 2, 1.4, 1, 1, 1, 1.8])
        _cols[0].markdown(_company or "—")
        _cols[1].markdown(_title or "—")
        _cols[2].markdown(_date or "—")
        # Colour-coded score badges
        def _score_badge(val: float) -> str:
            if val <= 0:
                return "—"
            col = "#22c55e" if val >= 70 else "#eab308" if val >= 50 else "#ef4444"
            return f'<span style="color:{col};font-weight:700;">{val:.0f}%</span>'
        _cols[3].markdown(_score_badge(_ats), unsafe_allow_html=True)
        _cols[4].markdown(_score_badge(_hr), unsafe_allow_html=True)
        _cols[5].markdown(_score_badge(_llm), unsafe_allow_html=True)

        # Inline status dropdown — updates on change
        _new_status = _cols[6].selectbox(
            label="",
            options=_TRACKER_STATUSES,
            index=_TRACKER_STATUSES.index(_status) if _status in _TRACKER_STATUSES else 0,
            key=f"tracker_status_{_entry_id}",
            label_visibility="collapsed",
        )
        if _new_status != _status:
            api(
                "PUT", f"/tracker/{_entry_id}",
                {"status": _new_status},
                token=st.session_state.token,
            )
            st.rerun()


# ─── Main router ─────────────────────────────────────────────────────────────

def render_footer():
    st.markdown("---")
    st.markdown(
        '<div style="text-align: center; padding: 16px 0; color: #64748b; font-size: 13px;">'
        'AI Resume Tuner &nbsp;|&nbsp; '
        '<a href="https://stats.uptimerobot.com/GrOgAbj0Nz" target="_blank" style="color: #818cf8; text-decoration: none;">System Status</a> &nbsp;|&nbsp; '
        '<a href="https://github.com/jananthan30/Resume-Builder" target="_blank" style="color: #818cf8; text-decoration: none;">GitHub</a> &nbsp;|&nbsp; '
        'MIT License'
        '</div>',
        unsafe_allow_html=True,
    )


def main():
    # Auto-load saved resume once per session when restored from URL param.
    # page_login() handles this on sign-in, but a page refresh skips login
    # so we need to re-fetch here whenever the session is authenticated but
    # the resume hasn't been loaded yet.
    if (is_authenticated()
            and not st.session_state.get("stored_resume")
            and not st.session_state.get("_resume_checked")):
        st.session_state._resume_checked = True
        saved = _fetch_saved_resume(st.session_state.token)
        if saved:
            st.session_state.stored_resume = saved
            st.session_state.resume_on_file = True

    render_nav()

    page = st.session_state.page

    if page == "home":
        page_home()
    elif page == "scorer":
        page_scorer()
    elif page == "discover":
        page_discover()
    elif page == "cover_letter":
        page_cover_letter()
    elif page == "rewriter":
        page_rewriter()
    elif page == "tracker":
        page_tracker()
    elif page == "register":
        page_register()
    elif page == "login":
        page_login()
    elif page == "dashboard":
        page_dashboard()
    else:
        page_home()

    render_footer()


if __name__ == "__main__":
    main()
